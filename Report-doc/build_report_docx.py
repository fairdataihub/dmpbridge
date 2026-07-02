"""Build report-doc/project_report.docx from structured content."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Styles ────────────────────────────────────────────────────────────────────
styles = doc.styles

def _set_font(style, name="Calibri", size=11, bold=False, color=None):
    f = style.font
    f.name  = name
    f.size  = Pt(size)
    f.bold  = bold
    if color:
        f.color.rgb = RGBColor(*color)

_set_font(styles["Normal"],   size=11)
_set_font(styles["Heading 1"], size=16, bold=True,  color=(31, 73, 125))
_set_font(styles["Heading 2"], size=13, bold=True,  color=(47, 84, 150))
_set_font(styles["Heading 3"], size=11, bold=True,  color=(68, 114, 196))

# tweak spacing on headings
for h in ("Heading 1", "Heading 2", "Heading 3"):
    pf = styles[h].paragraph_format
    pf.space_before = Pt(14 if h == "Heading 1" else 10)
    pf.space_after  = Pt(4)

styles["Normal"].paragraph_format.space_after = Pt(6)

# ── Helpers ───────────────────────────────────────────────────────────────────

def h1(text):
    doc.add_heading(text, level=1)

def h2(text):
    doc.add_heading(text, level=2)

def h3(text):
    doc.add_heading(text, level=3)

def para(text, bold_prefix=None):
    """Add a normal paragraph. If bold_prefix is given, that prefix is bolded."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def code_block(text):
    """Monospace block styled as a shaded paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.4)
    # light grey shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F2F2F2")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(9)
    return p

def table_with_header(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run  = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        # dark blue background
        tc  = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "1F497D")
        tcPr.append(shd)
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)

    # data rows
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        fill = "EEF3FA" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd  = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  fill)
            tcPr.append(shd)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()  # spacing after table
    return t

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F497D")
    pBdr.append(bottom)
    pPr.append(pBdr)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(40)
title_p.paragraph_format.space_after  = Pt(6)
r = title_p.add_run("DMPBridge — Project Status Report")
r.bold = True
r.font.name  = "Calibri"
r.font.size  = Pt(24)
r.font.color.rgb = RGBColor(31, 73, 125)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.paragraph_format.space_after = Pt(4)
r = sub_p.add_run("PDF-to-Structure Labeling Pipeline for Data Management Plans")
r.font.name  = "Calibri"
r.font.size  = Pt(13)
r.font.color.rgb = RGBColor(89, 89, 89)

doc.add_paragraph()  # blank line

meta = [
    ("Date",    "2026-07-02"),
    ("Version", "0.1.0"),
    ("Author",  "Nahid Zeinali"),
    ("Status",  "Active development — 7 experiments completed"),
]
for label, value in meta:
    mp = doc.add_paragraph()
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mp.paragraph_format.space_after = Pt(2)
    r = mp.add_run(f"{label}: ")
    r.bold = True
    r.font.size = Pt(11)
    mp.add_run(value).font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
h1("1. Project Overview")
divider()
para(
    "DMPBridge (v0.1.0) is a Python pipeline that reads Data Management Plan (DMP) PDF "
    "documents and classifies every text block into one of five semantic labels. The output "
    "is a flat labeled JSON file and an optional hierarchical 'structured' JSON following "
    "the DMP Tool narrative schema."
)
para(
    "Core problem: DMP PDFs mix funder-written template text with researcher-written "
    "responses. Automated systems that process DMPs need to distinguish these reliably. "
    "DMPBridge treats this as a block-level multi-class classification task, using LLMs "
    "as the classifier."
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. LABEL SCHEMA
# ══════════════════════════════════════════════════════════════════════════════
h1("2. Label Schema")
divider()
para("Every text block extracted from a PDF is assigned exactly one of five labels:")

label_rows = [
    ("title",                "Single main document title; appears once per document"),
    ("section.title",        "Numbered or named section heading (e.g. '1. Data sharing')"),
    ("section.description",  "Funder template text — instructions to the author; uses 'should', 'must'"),
    ("question.text",        "Sub-question prompt inside a section; asks the author to address a specific topic"),
    ("answer.text",          "Researcher's actual written response — narrative paragraphs describing plans"),
]
table_with_header(["Label", "Meaning"], label_rows, col_widths=[1.6, 4.8])

para(
    "Key distinction: section.description is written by the funder; answer.text is written "
    "by the researcher. question.text is a specific sub-prompt inside a section, not the "
    "section header itself."
)

# ══════════════════════════════════════════════════════════════════════════════
# 3. PIPELINE ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
h1("3. Pipeline Architecture")
divider()
para("The pipeline runs in three sequential stages:")

bullet("Extraction — pdfplumber reads the PDF page by page, line by line. Each line becomes "
       "one block dict (12 fields: page, line_order, text, x0, top, x1, bottom, avg_font_size, "
       "font_names, is_bold, is_italic, label). text_cleaner removes duplicate characters "
       "caused by layered PDF rendering. Typical output: 60–90 blocks per DMP document.",
       bold_prefix="Stage 1 — ")

bullet("Classification — blocks are sent to an LLM strategy (batch / whole-doc / pdf-direct) "
       "with a structured system prompt and few-shot examples. The LLM returns a JSON array "
       "mapping block IDs to label strings.",
       bold_prefix="Stage 2 — ")

bullet("Output — results are saved as a flat labeled JSON and optionally converted to a "
       "hierarchical structured JSON following the DMP Tool narrative schema.",
       bold_prefix="Stage 3 — ")

doc.add_paragraph()

h2("3.1  Subpackage Map")

code_block(
    "dmpbridge/\n"
    "├── core/\n"
    "│   ├── config.py            PROVIDER, MODEL, HOST, BATCH_SIZE, API keys\n"
    "│   ├── pipeline.py          process_pdf() orchestration\n"
    "│   └── converter.py         to_structured() → DMP Tool JSON schema\n"
    "├── preprocess/\n"
    "│   ├── pdfplumber_reader.py extract_blocks(); pdfplumber line extraction\n"
    "│   ├── text_cleaner.py      clean_blocks(), deduplicate doubled chars\n"
    "│   └── page_images.py       optional per-page PNG export with bounding boxes\n"
    "├── strategies/\n"
    "│   ├── batch.py             BatchStrategy: sliding-window batch calls\n"
    "│   ├── wholedoc.py          WholeDocStrategy: single call with all blocks\n"
    "│   └── pdf_direct.py        PdfDirectStrategy: raw PDF bytes → Claude API\n"
    "├── models/\n"
    "│   ├── ollama.py            OllamaModel (local server, structured output)\n"
    "│   ├── anthropic.py         AnthropicModel (Anthropic messages API)\n"
    "│   ├── openai.py            OpenAIModel (OpenAI chat completions)\n"
    "│   └── gemini.py            GeminiModel (Google Gemini API)\n"
    "├── prompts/\n"
    "│   ├── system.py            shared SYSTEM_PROMPT with few-shot examples\n"
    "│   └── labels.py            LABELS tuple, OUTPUT_SCHEMA\n"
    "├── parsers/\n"
    "│   └── json_parser.py       parse_llm_json(): strips markdown fences\n"
    "├── evaluation/\n"
    "│   ├── evaluate.py          extract_gold, evaluate_sample, gold_metrics\n"
    "│   └── experiment.py        ExperimentConfig, Experiment, YAML runner\n"
    "├── cli/\n"
    "│   ├── main.py              dmpbridge (single PDF, batch strategy)\n"
    "│   └── wholedoc_cmd.py      dmpbridge-wholedoc\n"
    "└── utils/\n"
    "    ├── logger.py            get_logger, setup_logging\n"
    "    └── exceptions.py        DmpBridgeError, ConfigurationError"
)

# ══════════════════════════════════════════════════════════════════════════════
# 4. LABELING STRATEGIES
# ══════════════════════════════════════════════════════════════════════════════
h1("4. Labeling Strategies")
divider()
para(
    "Three strategies share the same label schema and evaluation framework but differ in "
    "how they send content to the model."
)

h2("4.1  Batch Strategy")
para(
    "The default strategy. Extracts line-level blocks via pdfplumber, then calls the LLM "
    "multiple times using a sliding window."
)
bullet("Extract ~70 line-level blocks from the PDF via pdfplumber.")
bullet("Divide blocks into windows of batch_size (default: 10) with a sliding context of "
       "3 already-labeled blocks prepended to each window.")
bullet("Make ceil(N / batch_size) LLM calls; merge labels back by block index.")
doc.add_paragraph()
bullet("Supports all four providers (Ollama, Anthropic, OpenAI, Gemini).", bold_prefix="Advantage: ")
bullet("Multiple API calls per document; window boundary effects possible.", bold_prefix="Disadvantage: ")

code_block("dmpbridge-experiment experiments/claude-opus-4-8-batch.yaml")

h2("4.2  Whole-Document Strategy")
para(
    "Sends the entire block list to the model in a single API call. The model sees the "
    "full document structure at once."
)
bullet("No boundary effects — full document context in one call.", bold_prefix="Advantage: ")
bullet("High token usage; requires max_tokens=16384; only Anthropic and Ollama supported.", bold_prefix="Disadvantage: ")
code_block("dmpbridge-experiment experiments/claude-opus-4-8-wholedoc.yaml")

h2("4.3  PDF-Direct Strategy")
para(
    "Sends the raw PDF bytes (base64-encoded) directly to Claude's document API. "
    "Claude simultaneously extracts and classifies without using pdfplumber at all."
)
bullet(
    "Produces paragraph-level blocks (~20 per document) rather than line-level blocks (~70). "
    "Bounding-box coordinates are not available."
)
bullet(
    "Cross-strategy evaluation uses gold-based accuracy (gold_metrics()) rather than block "
    "accuracy, because paragraph-level blocks cannot be compared index-by-index with "
    "line-level blocks."
)
bullet("Only the Anthropic provider is supported — relies on Claude's PDF vision capability.", bold_prefix="Constraint: ")
code_block("dmpbridge-pdf --model claude-opus-4-8")

# ══════════════════════════════════════════════════════════════════════════════
# 5. PROVIDERS AND MODELS
# ══════════════════════════════════════════════════════════════════════════════
h1("5. Supported Providers and Models")
divider()

provider_rows = [
    ("Ollama (local)",  "DMPBRIDGE_HOST\n(default: localhost:11434)", "llama3.1:8b,  llama3.3:70b"),
    ("Anthropic",       "ANTHROPIC_API_KEY",                          "claude-opus-4-8"),
    ("OpenAI",          "OPENAI_API_KEY",                             "Configured — no experiments yet"),
    ("Gemini",          "GEMINI_API_KEY",                             "Configured — no experiments yet"),
]
table_with_header(["Provider", "Key / Setting", "Models Used"], provider_rows,
                  col_widths=[1.3, 2.2, 2.9])

para("Configuration is set in dmpbridge/core/config.py or via environment variables / .env:")
code_block(
    "DMPBRIDGE_PROVIDER=anthropic\n"
    "DMPBRIDGE_MODEL=claude-opus-4-8\n"
    "ANTHROPIC_API_KEY=sk-ant-..."
)

# ══════════════════════════════════════════════════════════════════════════════
# 6. EXPERIMENTS REGISTRY
# ══════════════════════════════════════════════════════════════════════════════
h1("6. Experiments Registry")
divider()
para(
    "Seven experiments have been run across three models and three strategies. Each "
    "experiment is defined by a YAML config in experiments/ and writes output to "
    "data/output/labeled/{tag}/."
)

exp_rows = [
    ("M01", "llama3.1:8b",     "batch",      "Ollama",    "llama3.1-8b_batch"),
    ("M02", "llama3.1:8b",     "wholedoc",   "Ollama",    "llama3.1-8b_whole_doc"),
    ("M03", "llama3.3:70b",    "batch",      "Ollama",    "llama3.3-70b_batch"),
    ("M04", "llama3.3:70b",    "wholedoc",   "Ollama",    "llama3.3-70b_whole_doc"),
    ("M05", "claude-opus-4-8", "batch",      "Anthropic", "claude-opus-4-8_batch"),
    ("M06", "claude-opus-4-8", "wholedoc",   "Anthropic", "claude-opus-4-8_whole_doc"),
    ("M07", "claude-opus-4-8", "pdf_direct", "Anthropic", "claude-opus-4-8_pdf"),
]
table_with_header(["ID", "Model", "Strategy", "Provider", "Output Tag"], exp_rows,
                  col_widths=[0.5, 1.6, 1.2, 1.2, 2.0])

h2("6.1  Experiment YAML Files")
code_block(
    "experiments/\n"
    "├── llama3.1-8b-batch.yaml\n"
    "├── llama3.1-8b-wholedoc.yaml\n"
    "├── llama3.3-70b-batch.yaml\n"
    "├── llama3.3-70b-wholedoc.yaml\n"
    "├── claude-opus-4-8-batch.yaml\n"
    "├── claude-opus-4-8-wholedoc.yaml\n"
    "└── claude-opus-4-8-pdf-direct.yaml"
)
para("Replay any experiment:")
code_block(
    "dmpbridge-experiment experiments/claude-opus-4-8-batch.yaml\n"
    "dmpbridge-experiment experiments/claude-opus-4-8-batch.yaml --evaluate\n"
    "dmpbridge-experiment --list"
)

# ══════════════════════════════════════════════════════════════════════════════
# 7. EVALUATION FRAMEWORK
# ══════════════════════════════════════════════════════════════════════════════
h1("7. Evaluation Framework")
divider()

h2("7.1  Gold Standard")
bullet("10 sample PDFs in data/input/pdfs/")
bullet("Manual annotation files: data/input/ground_truth/sampleN_dmp.json")
bullet("Reference text files: data/input/ground_truth/reference_text/sampleN_reference.txt")
bullet("Total gold items across all 10 samples: 741 labeled spans")

h2("7.2  Matching Algorithm")
para(
    "evaluate_sample() uses token-level containment with a 0.75 threshold:"
)
bullet("Forward check — for each predicted block, find the best-matching gold item.")
bullet("Reverse check — for each gold item, mark it __missed__ if no predicted block covers it.")
para(
    "This handles the mismatch between line-level pdfplumber blocks and paragraph-level "
    "gold items."
)

h2("7.3  Two Accuracy Measures")
acc_rows = [
    ("Block accuracy",
     "correct_blocks / total_predicted_blocks",
     "Comparing batch vs whole-doc within a model"),
    ("Gold accuracy",
     "correctly_covered_gold / total_gold_items",
     "Cross-strategy comparison; used for M07 (pdf-direct)"),
]
table_with_header(["Measure", "Formula", "When to Use"], acc_rows,
                  col_widths=[1.5, 2.4, 2.5])

h2("7.4  Known Results")
result_rows = [
    ("M07 — Opus 4.8 PDF-Direct", "~95.1%",   "~99.5%",    "Paragraph-level; no pdfplumber"),
    ("M01–M06 (all others)",       "See notebooks", "—",    "Loaded from data/output/labeled/"),
]
table_with_header(
    ["Experiment", "Block Accuracy", "Gold Accuracy", "Notes"],
    result_rows,
    col_widths=[2.0, 1.4, 1.4, 1.6]
)

# ══════════════════════════════════════════════════════════════════════════════
# 8. DATA FOLDER STRUCTURE
# ══════════════════════════════════════════════════════════════════════════════
h1("8. Data Folder Structure")
divider()
code_block(
    "data/\n"
    "├── input/\n"
    "│   ├── pdfs/\n"
    "│   │   └── sample1.pdf … sample10.pdf          10 DMP documents\n"
    "│   └── ground_truth/\n"
    "│       ├── sample1_dmp.json … sample10_dmp.json  manual annotations\n"
    "│       └── reference_text/\n"
    "│           └── sampleN_reference.txt             reference text for extraction eval\n"
    "└── output/\n"
    "    ├── extracted/\n"
    "    │   └── sample1.json … sample10.json          pdfplumber blocks, no labels\n"
    "    └── labeled/\n"
    "        ├── llama3.1-8b_batch/\n"
    "        ├── llama3.1-8b_whole_doc/\n"
    "        ├── llama3.3-70b_batch/\n"
    "        ├── llama3.3-70b_whole_doc/\n"
    "        ├── claude-opus-4-8_batch/\n"
    "        ├── claude-opus-4-8_whole_doc/\n"
    "        └── claude-opus-4-8_pdf/\n"
    "            ├── sample1.json … sample10.json\n"
    "            └── sample1_structured.json … sample10_structured.json"
)

# ══════════════════════════════════════════════════════════════════════════════
# 9. NOTEBOOKS
# ══════════════════════════════════════════════════════════════════════════════
h1("9. Notebooks")
divider()
para(
    "Seven numbered notebooks in notebooks/ cover the full pipeline from raw PDF to "
    "cross-model comparison. All notebooks verified to execute cleanly with "
    "jupyter nbconvert --execute."
)

nb_rows = [
    ("000_experiment_log.ipynb",
     "Experiment registry, strategy descriptions, cross-model summary"),
    ("001_pdfplumber_extraction.ipynb",
     "Run pdfplumber on all 10 PDFs; inspect block counts, font sizes, bold/italic patterns"),
    ("002_pdfplumber_extraction_eval.ipynb",
     "Compare extracted text vs reference text; word capture, ROUGE-L, F1, precision, recall"),
    ("003_strategy_comparison.ipynb",
     "Side-by-side accuracy across all 7 experiments; confusion matrices; cross-model agreement"),
    ("004_llama3.1-8b.ipynb",
     "Deep-dive evaluation for llama3.1:8b (batch + whole-doc)"),
    ("005_llama3.3-70b.ipynb",
     "Deep-dive evaluation for llama3.3:70b (batch + whole-doc)"),
    ("006_claude-opus-4-8.ipynb",
     "Deep-dive evaluation for claude-opus-4-8 (batch + whole-doc + pdf-direct)"),
]
table_with_header(["Notebook", "Purpose"], nb_rows, col_widths=[2.4, 4.0])

para(
    "Note: Notebook 003 loads all 7 variants. PDF-direct is excluded from Section 7 "
    "(cross-model block alignment) because paragraph-level blocks cannot be compared "
    "index-by-index with line-level blocks."
)

# ══════════════════════════════════════════════════════════════════════════════
# 10. CLI REFERENCE
# ══════════════════════════════════════════════════════════════════════════════
h1("10. CLI Reference")
divider()

cli_rows = [
    ("dmpbridge document.pdf",
     "batch",
     "Single PDF, any provider"),
    ("dmpbridge-wholedoc document.pdf",
     "wholedoc",
     "Single PDF, Anthropic or Ollama"),
    ("dmpbridge-pdf --model claude-opus-4-8",
     "pdf_direct",
     "All PDFs in data/input/pdfs/, Anthropic only"),
    ("dmpbridge-experiment experiments/X.yaml",
     "any",
     "Run full experiment from YAML config"),
    ("dmpbridge-experiment experiments/X.yaml --evaluate",
     "any",
     "Run experiment and print accuracy summary"),
    ("dmpbridge-experiment --list",
     "—",
     "List all available experiment configs"),
    ("dmpbridge-evaluate",
     "—",
     "Evaluate all samples for the default tag"),
    ("dmpbridge-evaluate path/to/sample1.json",
     "—",
     "Evaluate a single output file"),
]
table_with_header(["Command", "Strategy", "Description"], cli_rows,
                  col_widths=[2.8, 1.1, 2.5])

# ══════════════════════════════════════════════════════════════════════════════
# 11. SYSTEM PROMPT DESIGN
# ══════════════════════════════════════════════════════════════════════════════
h1("11. System Prompt Design")
divider()
para(
    "All three strategies share the same system prompt (dmpbridge/prompts/system.py). "
    "It uses four components:"
)
bullet("Clear label definitions — each label has a one-line description with key signal "
       "words ('should'/'must' → section.description; narrative text → answer.text).",
       bold_prefix="1. ")
bullet("Key distinctions — explicit rules for the three pairs most likely to be confused: "
       "section.description vs answer.text, question.text vs section.description.",
       bold_prefix="2. ")
bullet("Few-shot examples — 10+ real DMP excerpts covering every label, drawn from the "
       "actual 10-sample corpus.",
       bold_prefix="3. ")
bullet("Output constraint — 'You MUST output a JSON array with one entry for EVERY block "
       "— no explanation, no markdown.'",
       bold_prefix="4. ")

# ══════════════════════════════════════════════════════════════════════════════
# 12. KNOWN ISSUES
# ══════════════════════════════════════════════════════════════════════════════
h1("12. Known Issues and Limitations")
divider()

issue_rows = [
    ("Duplicate character rendering",
     "Blocks contain 'HHeelllloo' patterns",
     "Fixed in text_cleaner._deduplicate_chars()"),
    ("pdf-direct: no bounding boxes",
     "Cannot render visual overlays for M07",
     "By design; use gold accuracy for eval"),
    ("pdf-direct excluded from cross-model alignment",
     "Cannot compare block-by-block in Section 7",
     "Documented; gold accuracy used instead"),
    ("OpenAI and Gemini not benchmarked",
     "No accuracy data for these providers",
     "Providers implemented; YAML configs not yet created"),
    ("Whole-doc: large context requirement",
     "Fails on very long DMPs",
     "max_tokens=16384 set in WholeDocStrategy"),
]
table_with_header(["Issue", "Impact", "Status"], issue_rows,
                  col_widths=[1.9, 2.0, 2.5])

# ══════════════════════════════════════════════════════════════════════════════
# 13. DEPENDENCIES
# ══════════════════════════════════════════════════════════════════════════════
h1("13. Dependencies")
divider()

h2("Core (always required)")
for dep in [
    "pdfplumber ≥ 0.11 — PDF text and layout extraction",
    "requests ≥ 2.28 — HTTP client for Ollama",
    "Pillow ≥ 11.0 — optional per-page PNG export",
    "python-dotenv ≥ 1.0 — .env file support",
    "pyyaml ≥ 6.0 — experiment YAML configs",
]:
    bullet(dep)

h2("Optional — Notebooks")
for dep in ["matplotlib", "seaborn", "pandas"]:
    bullet(dep)

h2("Optional — REST Server")
for dep in ["fastapi", "uvicorn", "python-multipart"]:
    bullet(dep)

h2("Provider SDKs (install separately)")
for dep in [
    "anthropic — for Anthropic/Claude provider",
    "openai — for OpenAI provider",
    "google-generativeai — for Gemini provider",
]:
    bullet(dep)

# ══════════════════════════════════════════════════════════════════════════════
# 14. REPRODUCTION STEPS
# ══════════════════════════════════════════════════════════════════════════════
h1("14. Reproduction Steps")
divider()
code_block(
    "# 1. Install\n"
    "pip install -e \".[notebooks]\"\n\n"
    "# 2. Set API key (for Anthropic experiments)\n"
    "echo \"ANTHROPIC_API_KEY=sk-ant-...\" > .env\n\n"
    "# 3. Extract pdfplumber blocks (run notebook 001, or manually):\n"
    "#    Output → data/output/extracted/sampleN.json\n\n"
    "# 4. Run any experiment\n"
    "dmpbridge-experiment experiments/llama3.3-70b-batch.yaml\n"
    "dmpbridge-experiment experiments/claude-opus-4-8-batch.yaml --evaluate\n\n"
    "# 5. Run pdf-direct experiment\n"
    "dmpbridge-pdf --model claude-opus-4-8\n\n"
    "# 6. Open comparison notebook\n"
    "jupyter lab notebooks/003_strategy_comparison.ipynb"
)

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
out = "c:/Users/Nahid/dmpbridge/report-doc/project_report.docx"
doc.save(out)
print(f"Saved: {out}")
