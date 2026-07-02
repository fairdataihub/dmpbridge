"""
Update 000_experiment_log.ipynb (fix duplicate M08 + add Vision Batch strategy section)
and regenerate report-doc/project_report.md + project_report.docx.
"""
import json
from pathlib import Path

# ── 1.  Fix 000_experiment_log.ipynb ─────────────────────────────────────────
nb_path = Path("c:/Users/Nahid/dmpbridge/notebooks/000_experiment_log.ipynb")
nb = json.loads(nb_path.read_bytes().lstrip(b"\xef\xbb\xbf"))
cells = nb["cells"]

# Find cells by searching source content
def find_cell(keyword):
    for i, c in enumerate(cells):
        src = "".join(c.get("source", []))
        if keyword in src:
            return i, src
    return None, None

# ── Fix strategies cell: add Vision Batch ────────────────────────────────────
strat_i, strat_src = find_cell("### PDF-direct")
if strat_i is not None and "### Vision Batch" not in strat_src:
    vision_section = (
        "\n\n### Vision Batch\n\n"
        "Each page of the document is rendered as a 150 DPI PNG and sent to the model "
        "alongside the pdfplumber line-level blocks extracted from that page. "
        "Claude uses the visual layout (font size hierarchy, spacing, bold/italic rendering) "
        "together with the structured text to classify each block.\n\n"
        "- Context per call: one page image + pdfplumber blocks for that page\n"
        "- API calls per document: one per page (typically 1–3)\n"
        "- Provider: Anthropic API only (requires vision capability)\n"
        "- Block granularity: line-level (same as batch/whole-doc — fully comparable)\n"
        "- Page images saved to disk: `data/output/page_images/{stem}/page_NNN.png` "
        "(reusable across models)\n"
        "- Inference: `dmpbridge-experiment experiments/claude-opus-4-8-vision.yaml`"
    )
    new_src = strat_src.rstrip() + vision_section
    cells[strat_i]["source"] = new_src
    print(f"[000] Added Vision Batch to strategies cell (cell {strat_i})")

# ── Fix claude cell: remove malformed first M08 block ────────────────────────
claude_i, claude_src = find_cell("### M08")
if claude_i is not None:
    # The malformed block uses \path\ instead of `path` — find and remove it.
    # Strategy: split on "### M08" — keep everything before the first occurrence,
    # then locate the second clean occurrence and keep everything from there.
    marker = "### M08"
    first = claude_src.find(marker)
    second = claude_src.find(marker, first + 1)
    if second != -1:
        # Keep up to (not including) first M08, then append from the second M08 onwards
        new_src = claude_src[:first] + claude_src[second:]
        cells[claude_i]["source"] = new_src
        print(f"[000] Removed duplicate M08 section from cell {claude_i}")
    else:
        print(f"[000] Only one M08 found — no duplicate to remove")

nb_path.write_text(json.dumps(nb, ensure_ascii=False, indent=1), encoding="utf-8")
print("[000] Saved 000_experiment_log.ipynb")

# ── 2.  Recreate report-doc/project_report.md ────────────────────────────────
report_dir = Path("c:/Users/Nahid/dmpbridge/report-doc")
report_dir.mkdir(parents=True, exist_ok=True)
md_path = report_dir / "project_report.md"

md = """# DMPbridge — Project Report

**Date:** July 2026
**Dataset:** 10 manually labeled DMP documents, 741 text blocks
**Labels:** `title` · `section.title` · `section.description` · `question.text` · `answer.text`

---

## 1. Project Overview

**DMPbridge** is a Python pipeline for automatically labeling text blocks in Data Management Plan (DMP) PDF documents. The system extracts structured text blocks from PDFs and uses large language models to classify each block with one of five structural labels. The primary motivation is to support downstream tasks such as DMP comparison, funder-specific analysis, and automated compliance checking.

The pipeline supports multiple extraction strategies and multiple model backends, enabling systematic comparison of cost, accuracy, and block granularity across different approaches.

---

## 2. Label Schema

| Label | Description |
|---|---|
| `title` | Document title block |
| `section.title` | Section or sub-section heading |
| `section.description` | Funder-provided context or instructions for a section |
| `question.text` | A specific data management requirement item |
| `answer.text` | Researcher response or description |

**Key challenge:** `question.text` and `section.description` are the hardest pair to distinguish. Both are funder-written text; the boundary is structural (position in the DMP form) rather than lexical (word choice).

---

## 3. Architecture

### 3.1 Core Pipeline

```
PDF input
    │
    ├── pdfplumber extraction ──► text blocks (line-level, with bbox + font metadata)
    │       │
    │       ├── Batch strategy        (sliding 10-block windows)
    │       ├── Whole-doc strategy    (full document in one call)
    │       └── Vision Batch strategy (per-page PNG + blocks → Claude multimodal)
    │
    └── Raw PDF bytes ──► PDF-direct strategy (vision API, paragraph-level)

LLM classification
    │
    └── Labeled JSON output  ──► data/output/labeled/{tag}/sampleN.json
                                 structured JSON  ──► sampleN_structured.json
```

### 3.2 Module Layout

```
dmpbridge/
├── strategies/
│   ├── batch.py           — sliding-window batch strategy
│   ├── wholedoc.py        — whole-document strategy
│   ├── pdf_direct.py      — raw PDF bytes → Claude vision
│   ├── vision_batch.py    — per-page PNG + pdfplumber blocks
│   └── __init__.py        — get_strategy() factory
├── preprocess/
│   ├── pdfplumber_reader.py — extract_blocks()
│   └── page_images.py       — render_pages(), save_page_images()
├── evaluation/
│   ├── evaluate.py          — gold evaluation, F1 computation
│   └── experiment.py        — ExperimentRunner CLI
├── prompts/
│   └── default.py           — SYSTEM_PROMPT, LABELS
└── parsers/
    └── llm_json.py          — parse_llm_json() (handles markdown fences)
```

---

## 4. Prompting Strategies

### 4.1 Batch
Text blocks are classified in overlapping windows of 10 blocks with a 3-block sliding context overlap across windows. One API call per window (≈ total blocks / 7 calls per document). Preserves label continuity across window boundaries.

### 4.2 Whole-document
All extracted blocks from the document are passed to the model in a single API call. Full visibility into document structure. Uses `max_tokens=16384` for Claude to accommodate full structured output.

### 4.3 PDF-direct
Raw PDF bytes are sent directly to a vision-capable model (Claude). No pdfplumber extraction. The model reads the PDF and returns paragraph-level blocks in a single call. Produces fewer but larger blocks (~20 per document vs ~74 for pdfplumber). No bounding-box coordinates.

### 4.4 Vision Batch (new — M08)
pdfplumber extracts line-level blocks from the PDF. Each page is rendered as a 150 DPI PNG and base64-encoded. Claude receives the page image alongside the blocks extracted from that page in a single multimodal API call. The model uses the visual layout (font size hierarchy, bold/italic rendering, spacing) to assist classification.

Key design decisions:
- Page images saved to disk (`data/output/page_images/{stem}/page_NNN.png`) and reused across experiments
- Page-local block IDs (0, 1, 2…N per page) sent per call; mapped back to global block list after response
- Block granularity identical to batch/whole-doc — fully comparable metrics
- Anthropic-only (requires vision API capability)

---

## 5. Experiment Registry

| ID | Model | Strategy | Output Tag | Samples | Accuracy | Status |
|---|---|---|---|---|---|---|
| M01 | Claude Opus 4.8 | Batch | `claude-opus-4-8_batch` | 10/10 | 94.9% | Done |
| M02 | Claude Opus 4.8 | Whole-doc | `claude-opus-4-8_whole_doc` | 10/10 | 96.9% | Done |
| M03 | Llama 3.3 70B | Batch | `llama3.3-70b_batch` | 10/10 | 94.1% | Done |
| M04 | Llama 3.3 70B | Whole-doc | `llama3.3-70b_whole_doc` | 10/10 | 91.8% | Done |
| M05 | Llama 3.1 8B | Batch | `llama3.1-8b_batch` | 10/10 | 67.3% | Done |
| M06 | Llama 3.1 8B | Whole-doc | `llama3.1-8b_whole_doc` | 10/10 | 84.2% | Done |
| M07 | Claude Opus 4.8 | PDF-direct | `claude-opus-4-8_pdf` | 10/10 | 95.1% / 99.5%† | Done |
| M08 | Claude Opus 4.8 | Vision Batch | `claude-opus-4-8_vision` | 10/10 | 94.7% | Done |
| F01 | Llama 3.1 8B fine-tune | — | — | — | — | Planned |

† M07 PDF-direct: 95.1% block-level (193/203 paragraph blocks), 99.5% gold-based (correct / 741 gold items).

---

## 6. Results Summary

### 6.1 Claude Opus 4.8 — All Strategies

| Strategy | Total Blocks | Correct | Accuracy | Notes |
|---|---|---|---|---|
| Batch (M01) | 741 | 703 | 94.9% | Baseline |
| Whole-doc (M02) | 741 | 718 | 96.9% | +2.0pp vs batch |
| PDF-direct (M07) | 203 / 741† | 193 / 737† | 95.1% / 99.5%† | No bbox |
| Vision Batch (M08) | 741 | 702 | 94.7% | Line-level + visual |

### 6.2 Cross-model Comparison (Batch Strategy)

| Model | Batch | Whole-doc | Δ Whole−Batch |
|---|---|---|---|
| Claude Opus 4.8 | 94.9% | 96.9% | +2.0pp |
| Llama 3.3 70B | 94.1% | 91.8% | −2.3pp |
| Llama 3.1 8B | 67.3% | 84.2% | +16.9pp |

### 6.3 Key F1 Findings

**question.text F1** — the hardest label:

| Variant | q.text F1 | sec.desc F1 |
|---|---|---|
| Claude Opus 4.8 — Whole-doc | 88.2% | 99.1% |
| Claude Opus 4.8 — PDF-direct | 81.6% | 100.0% |
| Claude Opus 4.8 — Batch | 62.0% | 84.4% |
| Claude Opus 4.8 — Vision Batch | 60.9% | 83.3% |
| Llama 3.3 70B — Batch | 70.7% | 87.7% |
| Llama 3.3 70B — Whole-doc | 29.0% | 78.6% |
| Llama 3.1 8B — Whole-doc | 47.9% | 74.2% |
| Llama 3.1 8B — Batch | 26.0% | 41.0% |

---

## 7. Strategy Trade-offs

| Strategy | Block granularity | Bbox? | Visual context | API calls/doc | Best for |
|---|---|---|---|---|---|
| Batch | Line-level | Yes | No | ~10 | Local models, cost control |
| Whole-doc | Line-level | Yes | No | 1 | Claude — highest accuracy |
| PDF-direct | Paragraph-level | No | Yes (PDF) | 1 | Maximum gold coverage |
| Vision Batch | Line-level | Yes | Yes (per page) | 1–3 | Research; reusable images |

**Recommendation by use case:**
- Highest accuracy with bounding boxes: **Whole-doc** (96.9%)
- Highest gold label coverage: **PDF-direct** (99.5% gold-based)
- Best offline / no-API option: **Llama 3.3 70B Batch** (94.1%, free)
- Balanced with visual context + line precision: **Vision Batch** (94.7%)

---

## 8. Evaluation Methodology

### 8.1 Block-level Accuracy
For pdfplumber strategies (batch, whole-doc, vision batch):
```
accuracy = correct_blocks / total_blocks
```
where a block is "correct" if the predicted label exactly matches the manually assigned gold label.

### 8.2 Gold-based Accuracy (PDF-direct)
PDF-direct produces paragraph-level blocks that do not align 1:1 with pdfplumber blocks. A separate gold-based metric is used:
```
gold_accuracy = matched_correct / total_gold_items
```
Each predicted paragraph block is matched to the nearest gold item by text overlap.

### 8.3 Per-label F1
Computed across all 10 samples combined (micro-average per label):
```
F1 = 2 * precision * recall / (precision + recall)
```

---

## 9. Data

### 9.1 Input
- `data/input/pdfs/sample1.pdf` … `sample10.pdf` — 10 DMP PDFs
- `data/input/ground_truth/sample1_dmp.json` … — manually labeled gold blocks

### 9.2 Output
- `data/output/labeled/{tag}/sampleN.json` — predicted labels per block
- `data/output/labeled/{tag}/sampleN_structured.json` — structured section hierarchy
- `data/output/page_images/sampleN/page_NNN.png` — page PNGs (150 DPI, shared across experiments)

### 9.3 Block Statistics
| Statistic | Value |
|---|---|
| Total gold blocks | 741 |
| Average blocks per document | 74.1 |
| Average pages per document | 1.8 |
| PDF-direct block count | 203 (~20/doc) |

---

## 10. Evaluation Notebooks

| Notebook | Covers |
|---|---|
| `000_experiment_log.ipynb` | Registry, strategy descriptions, cross-model summary |
| `003_strategy_comparison.ipynb` | All-variant accuracy, F1, block-level analysis, cross-model agreement |
| `004_llama3.1-8b.ipynb` | Llama 3.1 8B deep dive (M05, M06) |
| `005_llama3.3-70b.ipynb` | Llama 3.3 70B deep dive (M03, M04) |
| `006_claude-opus-4-8.ipynb` | Claude Opus 4.8 deep dive (M01, M02, M07, M08) |

---

## 11. Reproduction

```bash
# Activate environment
.venv\\Scripts\\activate   # Windows
source .venv/bin/activate  # macOS/Linux

# Run an experiment
dmpbridge-experiment experiments/claude-opus-4-8-vision.yaml

# PDF-direct (Claude only)
dmpbridge-pdf --model claude-opus-4-8

# Execute all evaluation notebooks
jupyter nbconvert --to notebook --execute --inplace notebooks/*.ipynb
```

Environment variables required:
- `ANTHROPIC_API_KEY` — for Claude experiments (M01, M02, M07, M08)
- No key required for Ollama experiments (M03–M06); Ollama must be running locally

---

## 12. Future Work

### F01 — Fine-tune Llama 3.1 8B
**Hypothesis:** Fine-tuning on labeled DMP blocks will close most of the accuracy gap between Llama 3.1 8B (84.2% whole-doc) and the 70B model (94.1% batch), making low-resource deployment viable.

**Blocker:** The current 10-sample dataset is too small to split into train/eval sets. Dataset expansion (target: 30+ samples) is a prerequisite.

**Planned approach:** LoRA or QLoRA fine-tuning; evaluate on a held-out set; target `question.text` F1 ≥ 70%.

### Other Possible Work
- Sequence correction post-processing (state machine over label sequence to enforce structural constraints)
- Dataset expansion to non-NIH/NSF funders
- Multi-model ensemble (majority vote across Claude, Llama 3.3 70B)
"""

md_path.write_text(md, encoding="utf-8")
print("[report] Wrote project_report.md")

# ── 3.  Regenerate project_report.docx ───────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    import re

    doc = Document()

    # Styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def set_heading(para, level, text):
        para.style = doc.styles[f"Heading {level}"]
        para.runs[0].text = text if para.runs else None
        if not para.runs:
            run = para.add_run(text)

    def add_heading(doc, text, level):
        p = doc.add_heading(text, level=level)
        return p

    def add_code(doc, text):
        p = doc.add_paragraph(style="No Spacing")
        run = p.add_run(text)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        p.paragraph_format.left_indent = Inches(0.3)
        return p

    def add_table_from_md(doc, md_lines):
        rows = [l for l in md_lines if l.strip().startswith("|") and "---" not in l]
        if not rows:
            return
        def parse_row(line):
            return [c.strip() for c in line.strip().strip("|").split("|")]
        headers = parse_row(rows[0])
        data    = [parse_row(r) for r in rows[1:]]
        t = doc.add_table(rows=1 + len(data), cols=len(headers))
        t.style = "Table Grid"
        hrow = t.rows[0]
        for i, h in enumerate(headers):
            cell = hrow.cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            shading = cell._tc.get_or_add_tcPr()
        for ri, row_data in enumerate(data):
            row = t.rows[ri + 1]
            for ci, val in enumerate(row_data):
                row.cells[ci].text = val

    # Parse markdown and build docx
    lines = md.split("\n")
    i = 0
    code_buf = []
    table_buf = []
    in_code = False
    in_table = False

    while i < len(lines):
        line = lines[i]

        # Code block
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                for cl in code_buf:
                    add_code(doc, cl)
                code_buf = []
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # Table
        if line.strip().startswith("|"):
            table_buf.append(line)
            i += 1
            continue
        elif table_buf:
            add_table_from_md(doc, table_buf)
            table_buf = []
            doc.add_paragraph()

        # Headings
        if line.startswith("# ") and not line.startswith("## "):
            add_heading(doc, line[2:], 1)
        elif line.startswith("## "):
            add_heading(doc, line[3:], 2)
        elif line.startswith("### "):
            add_heading(doc, line[4:], 3)
        elif line.startswith("#### "):
            add_heading(doc, line[5:], 4)
        elif line.strip() == "---":
            doc.add_paragraph()
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(re.sub(r"`([^`]+)`", r"\1", line[2:]))
        elif line.strip():
            p = doc.add_paragraph()
            # Inline code: strip backtick markers for docx
            text = re.sub(r"`([^`]+)`", r"\1", line)
            text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
            p.add_run(text)
        else:
            if i > 0 and lines[i-1].strip():
                doc.add_paragraph()
        i += 1

    if table_buf:
        add_table_from_md(doc, table_buf)

    docx_path = report_dir / "project_report.docx"
    doc.save(str(docx_path))
    print(f"[report] Wrote project_report.docx ({docx_path.stat().st_size:,} bytes)")

except Exception as e:
    print(f"[report] docx generation failed: {e}")
    print("[report] project_report.md was still written — generate docx manually with pandoc")

print("\nAll done.")
