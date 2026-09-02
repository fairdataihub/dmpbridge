"""Build Report-doc/design_document.docx — "DMPbridge, Pipeline Strategy".

A design document following the outline supplied on 28 August 2026: overview,
literature review, label taxonomy, annotation strategy, extraction,
classification, prompting, evaluation, results. Every section is filled from
what the project has on disk (README, prompt, worklogs, evaluation code, the
current run results); wherever something is not in the project — literature,
annotator details, funder provenance — a highlighted "What you need to add"
box says exactly what is missing rather than inventing it.

Numbers quoted are the runs on disk as of 28 August 2026 (gemma4:e4b ×
three extractors, four models × pdfplumber), 75% containment, both paths.
Re-run this script after any re-run so they do not go stale.

    python scripts/build/build_design_doc.py
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

TITLE_BLUE = RGBColor(0x1E, 0x40, 0x6E)
SUB_BLUE   = RGBColor(0x22, 0x63, 0x99)
H2_BLUE    = RGBColor(0x4F, 0x81, 0xBD)
GREY       = RGBColor(0x59, 0x59, 0x59)
RUST       = RGBColor(0x99, 0x3A, 0x2B)

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(10.5)
for s in doc.sections:
    s.left_margin = s.right_margin = Inches(0.98)
    s.top_margin = s.bottom_margin = Inches(0.87)


def shade(cell, hexcolor):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def setc(cell, text, *, bold=False, size=9, color=None, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color


def table(headers, rows, widths=None, size=9, center_from=1):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]
        setc(c, hd, bold=True, size=size, color=RGBColor(0xFF, 0xFF, 0xFF), center=i >= center_from)
        shade(c, "1E406E")
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, v in enumerate(row):
            setc(cells[ci], str(v), bold=(ci == 0), size=size, center=ci >= center_from)
            if ri % 2 == 1:
                shade(cells[ci], "F2F5FA")
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = TITLE_BLUE
    return p


def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = H2_BLUE
    return p


def para(text, *, size=10.5, bold=False, italic=False, color=None, after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return p


def bullet(text, prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    if prefix:
        r = p.add_run(prefix)
        r.bold = True
        r.font.size = Pt(10.5)
    r = p.add_run(text)
    r.font.size = Pt(10.5)


def mono(lines, size=8.5):
    for ln in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(ln)
        r.font.name = "Consolas"
        r.font.size = Pt(size)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def todo(items, title="What you need to add"):
    """A highlighted box listing what the project does not yet contain."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    c = t.rows[0].cells[0]
    shade(c, "FDF4E9")
    c.text = ""
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RUST
    for it in items:
        p = c.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.15)
        r = p.add_run("• " + it)
        r.font.size = Pt(9.5)
        r.font.color.rgb = RUST
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ── Header ───────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run("DMPbridge")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = TITLE_BLUE
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(10)
r = p.add_run("Pipeline Strategy — design document")
r.font.size = Pt(13)
r.font.color.rgb = SUB_BLUE
for line in ["Date: 28 August 2026",
             "Status: draft — sections marked with a highlighted box need input that the project does not contain",
             "Companion: Report-doc/project_report.docx (engineering detail), Report-doc/worklog/ (day-by-day reasoning)"]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(line)
    r.bold = True
    r.font.size = Pt(10.5)
doc.add_paragraph()

# ── 1 ────────────────────────────────────────────────────────────────────
h1("1. Overview")
para("Data Management Plans (DMPs) are commonly required by funders with every grant proposal, but the "
     "formats and requirements imposed by each funder keep evolving, and no two funders' PDFs are structured "
     "the same way. DMPbridge is an open-source (MIT License), Python-based pipeline that converts DMP PDFs "
     "from any funder format into DMP Tool JSON, combining a narrative portion that mirrors DMP Tool's "
     "internal structure with the RDA DMP Common Standard JSON for machine-actionable output.")
para("The pipeline has two stages that matter and two that are bookkeeping. A PDF is first read into one "
     "text string per document with its visual formatting marked in the text (bold, italic, underline). A "
     "local large language model then reads that string and cuts it into labelled pieces — title, section "
     "heading, funder instruction, question, answer. The labelled pieces are nested into the DMP Tool "
     "narrative schema, and a small set of annotation rules fills in questions the source left implicit. "
     "Every stage is written to disk so any document can be followed through.")
para("Everything runs locally through Ollama: no API cost, and no document leaves the machine. The "
     "recommended configuration is pdfplumber for reading and Gemma 4 e4b for labelling; it scores 0.946 "
     "F1 on the ten-document evaluation set (section 9).", size=10, color=GREY)
todo(["Confirm the scope statement about RDA DMP Common Standard output: the current converter writes the "
      "DMP Tool narrative schema (narrative → template → section[] → question[] → answer); the RDA "
      "maDMP JSON is not produced by the code on disk. Either add it as planned work or reword.",
      "Add the intended users and the deployment picture (who runs this, on what, how often)."])

# ── 2 ────────────────────────────────────────────────────────────────────
h1("2. Literature Review")
para("This section is not written: the project contains no bibliography, and nothing below should be cited "
     "without being checked against the source. What the design draws on, and therefore what the review "
     "should cover:")
bullet("what DMP Tool and the RDA DMP Common Standard (maDMP) specify, and prior work on turning "
       "narrative DMPs into machine-actionable records.", prefix="Machine-actionable DMPs — ")
bullet("pdfplumber (rule-based, MIT); Docling (IBM, 2024, layout-model based, designed for downstream LLM "
       "use); LightOnOCR-2-1B (vision-language OCR model). Prior applications of each to document "
       "structure extraction.", prefix="PDF structure extraction — ")
bullet("classifying document segments with LLMs from a single prompt; the effect of visual-formatting cues; "
       "grammar-constrained JSON output.", prefix="LLM text-block classification — ")
bullet("token-overlap matching, precision/recall/F1 on segment labels, and evaluation against revised "
       "annotations.", prefix="Evaluation of document segmentation — ")
todo(["Write the review and add the references the outline numbers as 1 and 2–4 (pdfplumber in similar "
      "work; Docling's technical report and its prior uses).",
      "Cite the DMP Tool schema and the RDA DMP Common Standard specification versions targeted.",
      "If the Llama, Gemma and Qwen model cards are cited, note the exact tags used here: llama3.1:8b, "
      "llama3.3:70b, gemma4:e4b, qwen2.5:14b (Ollama)."])

# ── 3 ────────────────────────────────────────────────────────────────────
h1("3. Label Taxonomy")
para("Every piece of text in a DMP is given exactly one of five labels. The definitions below are the ones "
     "the model is given (dmpbridge/prompts/constants.py); the last column is what the annotation data "
     "actually shows, which is the reason two of the definitions were rewritten (section 7).")
table(["Label", "Meaning", "Written by", "What the data shows"],
      [["title", "The single main title of the document — appears once, typically at the top, usually short.",
        "Researcher", "One per document; largest or first line on page 1"],
       ["section.title", "A heading that opens a new top-level section. Often numbered (1., 2.) or a named "
        "phrase (\"Element 1:\").", "Funder / researcher",
        "43 in the set; numbered or 'Element N:' — never lettered"],
       ["section.description", "Text that explains what a section is about — funder instructions, not a "
        "question to the researcher and not the researcher's response. Typically follows a section.title.",
        "Funder", "8 in the set; usually italic instruction paragraphs"],
       ["question.text", "A specific question, instruction or prompt that asks the researcher to address a "
        "particular topic. Usually ends in a colon or reads as a direct ask.", "Funder",
        "16 in the set; lettered sub-items (A., B., C.) and short bold labels inline with their answer"],
       ["answer.text", "The researcher's actual written response — narrative paragraphs describing what the "
        "team will do, has done or plans to do, usually in first or third person about the team.",
        "Researcher", "55 in the set"]],
      widths=[1.15, 2.65, 0.95, 1.75], size=8.5, center_from=99)
para("The hard boundary is question.text against section.description: both are funder-written, and the "
     "difference is structural — a question asks, a description explains — rather than lexical. The second "
     "hard case is a bold label that shares a line with its answer (\"Roles & Responsibilities. For the "
     "proposed research…\"): the annotation splits it into a question and an answer, so the extractor's "
     "formatting markers are what make the split visible.", size=10, color=GREY)

# ── 4 ────────────────────────────────────────────────────────────────────
h1("4. Annotation Strategy")
para("We selected 10 DMPs, 23 pages in total (1–5 pages each), covering several funder templates — the set "
     "includes NIH-style plans organised by \"Element 1–6\" with lettered sub-questions, a DOE-style template "
     "whose instruction text is set in italics, and NSF-style plans with numbered sections. Each was "
     "annotated by hand into the DMP Tool narrative schema, giving 132 labelled items under the original "
     "annotation (10 titles, 43 section titles, 8 descriptions, 16 questions, 55 answers).")
para("Two annotation versions exist, and both are kept:")
bullet("the annotation as first produced; the model's output is scored against it directly (Path A).",
       prefix="Original — ")
bullet("a later pass in which questions the source leaves implicit were filled in from the nearest heading, "
       "raising the question count from 16 to 56. The filling was written down as a 16-row truth table "
       "(data/input/Rules.xlsx) over which of title, section.title, section.description and question.text "
       "are empty, and the pipeline applies that table mechanically after structuring; the result is scored "
       "against the revised annotation (Path B). The rules reproduce the revised annotation on 9 of the 10 "
       "documents; the tenth (sample 5) merges several sub-questions into one, a pattern with no derivable "
       "trigger.", prefix="Revised — ")
para("A plain-text reference transcription of each PDF (data/input/reference_text/) sits beside the "
     "annotations for checking extraction coverage.", size=10, color=GREY)
para("A trade-off that should be stated wherever the scores are quoted: samples 1 and 2 were used while "
     "writing the prompt, so figures that include them are optimistic and not a clean estimate on unseen "
     "documents. They were nonetheless kept in the evaluation set because they hold most of the question.text "
     "items — on the other eight documents that label had only six gold items and was unmeasurable.",
     size=10, italic=True, color=RUST)
todo(["Who annotated (how many people), what guidelines they followed, and whether agreement between "
      "annotators was measured.",
      "Where the ten PDFs came from — funder and programme per sample — and the selection criteria. The "
      "funder attributions above are inferred from the documents' own headings and should be confirmed.",
      "Why the annotation was revised, and who decided the fill-in rules in Rules.xlsx."])

# ── 5 ────────────────────────────────────────────────────────────────────
h1("5. Structured Text Extraction from PDF")
para("The goal of this stage is to convert the raw PDF into text that preserves the layout evidence a "
     "reader uses — which words are bold, italic or underlined, how big they are, where they sit — because "
     "those signals carry the document's structure. The design choice that shaped everything else: rather "
     "than a list of blocks with metadata fields, the extractor produces one string per document with the "
     "formatting marked inline (**bold**, _italic_, ++underline++), and the model reads that string whole. "
     "Extraction and segmentation are therefore not separate steps; the model decides the block boundaries.")
para("We tried pdfplumber (MIT), Docling (MIT) and LightOnOCR-2-1B. All are open-source. pdfplumber is a "
     "lightweight, rule-based library used in similar work before. Docling, developed by IBM, is newer (first "
     "release 2024), AI-based (a layout model, with OCR available), used in similar work and designed for "
     "downstream LLM analysis. LightOnOCR-2-1B is a one-billion-parameter vision-language model that "
     "transcribes a page image.")
table(["Extractor", "What it reads", "Where bold / italic / underline come from", "Cost"],
      [["pdfplumber", "the PDF's own characters and drawn shapes",
        "font name and size per character, judged against the document's most common (body) font; an "
        "underline is a thin rectangle drawn under the word — a signal fonts cannot carry", "CPU, < 1 s / doc"],
       ["Docling", "the PDF's characters, grouped into blocks by a layout model; the parsed page cells kept",
        "the same font rules applied to Docling's own word cells; hyperlink rectangles as underline; the "
        "layout model's heading label where the font marks nothing", "CPU, 0.1–3 s / doc"],
       ["LightOnOCR-2-1B", "a picture of each page (150 dpi)",
        "the model's own judgement while transcribing; it is asked for the markers and follows the request "
        "loosely (headings as #, few italics, never an underline)", "CUDA GPU, ~30–60 s / doc"]],
      widths=[1.05, 1.65, 2.85, 0.95], size=8.5, center_from=99)
para("Two findings from the extraction work decided the design:", size=10.5)
bullet("Docling's Markdown and JSON exports contain no bold, italic or underline — its PDF pipeline reads the "
       "fonts into its page cells and then discards them when it assembles the document. The first Docling "
       "extractor worked from the Markdown and gave the model headings only. The extractor was rebuilt to "
       "keep Docling's parsed pages and apply pdfplumber's rules to its word cells; the two extractors now "
       "produce identical markers on 8 of the 10 documents (sample 2: 42 bold, 41 italic, 11 underlined for "
       "both). That change alone was worth 16 points of F1.", prefix="Read the native data, not the export. ")
bullet("Sample 6's five headings are underlined and nothing else — same font and size as the body. Fonts "
       "cannot see that; pdfplumber's drawn rectangles can, and detecting them took that document from 0.52 "
       "to 1.00. Docling exposes no drawn shapes at any level, so this one signal is the whole remaining "
       "difference between the two extractors. LightOnOCR sees the underline in the image and reports it as "
       "bold, which is enough.", prefix="Underline is a shape, not a font attribute. ")
para("Coverage: pdfplumber and Docling recover the same words — a per-document probe found 0–1 words present "
     "in one extractor's text and absent from the other's. Both extractors can also write their raw reading "
     "of a PDF beside the pipeline text (sampleN.native.json: every word with its font and box, rectangles, "
     "hyperlinks; for Docling its full conversion result) so any marker can be traced to what produced it.",
     size=10, color=GREY)
todo(["Add the numbered references for pdfplumber's and Docling's prior use (the outline's citations 1 and "
      "2–4).",
      "If OCR-only (scanned) PDFs are in scope, say so: pdfplumber cannot read them, Docling falls back to "
      "RapidOCR, LightOnOCR reads them natively. None of the ten evaluation documents is scanned, so this is "
      "untested."])

# ── 6 ────────────────────────────────────────────────────────────────────
h1("6. Structural and Semantic Classification")
para("The goal of this stage is to classify the extracted text into the five labels using an LLM, "
     "leveraging both the structural evidence carried by the markers (bold, italic, underline, position) and "
     "the semantic content of the text. The model receives the whole document in one call and returns a flat "
     "list of {text, label} items; Ollama constrains generation to that JSON schema, so the output is always "
     "well-formed and no repair step exists. Decoding is deterministic (temperature 0.0, 32,768-token "
     "context); one configuration run three times produced identical counts, so the run-to-run noise floor is "
     "±0.002 F1 and differences above about 0.005 are real.")
para("We tested Llama 3.1 8B, Llama 3.3 70B, Gemma 4 e4b and Qwen 2.5 14B, all through Ollama. Llama offers "
     "the longest context window among strong open models; Gemma 4 e4b is new and small (3.4 GB of VRAM). "
     "Gemma 4 e4b is the clear best on this task and also the fastest: 0.946 F1 against 0.779 for the 70B "
     "model at roughly a twentieth of the runtime.")
table(["Model", "VRAM", "10 documents", "Path A F1 (pdfplumber)"],
      [["gemma4:e4b", "3.4 GB", "~2.5 min", "0.946"],
       ["llama3.3:70b", "~40 GB", "~40 min", "0.779"],
       ["qwen2.5:14b", "~9 GB", "~4 min", "0.748"],
       ["llama3.1:8b", "~5 GB", "~2.5 min", "0.678"]],
      widths=[1.4, 1.0, 1.3, 1.7])
para("Two properties of the smaller models shaped the prompt work: prompt changes do not transfer between "
     "models (one wording gained 0.6 points on llama3.1:8b and cost 4.0 on gemma4:e4b), and llama3.1:8b does "
     "not act on negative instructions — told that lettered sub-items are not section titles, it labels "
     "seven of eight as section titles regardless.", size=10, color=GREY)
todo(["The VRAM figures for the 70B, 14B and 8B models above are approximate; replace with the ollama ps "
      "readings if they are to be quoted."])

# ── 7 ────────────────────────────────────────────────────────────────────
h1("7. Prompting Strategy")
para("A single system prompt is used for every model and extractor. It contains, in order:")
bullet("the five labels, one sentence each (section 3).", prefix="Definitions — ")
bullet("what the markers mean and how to use them: an emphasised short phrase followed by plain text on the "
       "same line is a label to be split from its answer; longer emphasised or italic passages that explain "
       "rather than ask are descriptions; not all emphasis is structure, and an underlined link is not a "
       "heading; absence of markers does not rule a label out.", prefix="Formatting markers — ")
bullet("process the whole document, skip nothing; reproduce text verbatim without the markers; do not invent "
       "placeholder answers; do not force a question where none is asked.", prefix="Rules — ")
bullet("one short example per label, drawn from samples 1 and 2.", prefix="Examples — ")
bullet("the JSON array of {text, label}, enforced by schema.", prefix="Output format — ")
para("The extracted document is sent separately as the user message, so the prompt is not a template and "
     "contains no placeholder. Of the two approaches originally planned — the whole document in one prompt, "
     "or block by block to reduce context — only the whole-document one is implemented. The longest document "
     "in the set is 15,000 characters, far below the 32k-token context, so block-by-block has not been "
     "needed; it has also never been compared against.")
para("Three things learned about the prompt, each measured:", size=10.5)
bullet("moving three blank lines with no wording change moved F1 by 1.9 points and false positives by 13 on "
       "llama3.1:8b, reproducibly. The prompt file is treated as whitespace-significant.",
       prefix="It is whitespace-sensitive. ")
bullet("the section.title definition originally said lettered sub-items are section titles; the data says "
       "they are always questions. Correcting it fixed gemma4:e4b completely (5 of 8 mislabelled → 0) and "
       "llama3.1:8b not at all.", prefix="Its wording was measurably wrong once. ")
bullet("every result must come from a single prompt version, and every model must be re-run after any "
       "change; a partial re-run produces figures that cannot be compared.",
       prefix="Changes do not transfer between models. ")
todo(["The prompt was edited on 25 August (one sentence added; a doubled bullet dash that may be a typo). "
      "The Docling results were produced under that version; the pdfplumber, LightOnOCR and other-model "
      "results under the previous one. Re-run gemma4:e4b with pdfplumber and LightOnOCR under the current "
      "prompt before quoting the extractor comparison as a single-prompt result.",
      "State whether the few-shot examples taken from samples 1 and 2 will be replaced before the final "
      "evaluation, since those two documents are in the evaluation set."])

# ── 8 ────────────────────────────────────────────────────────────────────
h1("8. Evaluation Strategy")
para("Predictions and ground truth are both flattened to (text, label) pairs from the structured JSON. Each "
     "gold item claims the best unclaimed predicted item whose words it contains at least 75% of (partial "
     "credit; the project ran exact-match briefly and reverted). A predicted item claimed by nothing is a "
     "false positive; a gold item claiming nothing is a false negative; a claimed item with the wrong label is "
     "both. One matching routine feeds the metrics, the confusion matrix and the error listings.")
para("We evaluate every configuration on two paths, and always report both:")
table(["", "What is scored", "Against"],
      [["Path A", "stage 3 — the model's structured output as produced", "the original annotation (132 items)"],
       ["Path B", "stage 4 — the same output after the Rules.xlsx fill-in", "the revised annotation (174 items)"]],
      widths=[0.8, 3.0, 2.6], center_from=99)
para("The gap between the two paths measures the rules' contribution; it is not a second opinion on one "
     "number, and the paths use different item counts, so scores are compared, never counts.",
     size=10, color=GREY)
para("Metrics: micro-averaged precision, recall and F1 over all items (used for ranking), per-label "
     "precision/recall/F1, and a confusion matrix whose last column is gold items the model never produced. "
     "Every configuration is also scored per document, so one document cannot hide inside an average.")
para("We also compared the extractors on text-level measures before any model ran — words captured "
     "(pdfplumber and Docling differ by 0–1 words per document) and marker counts per document against "
     "pdfplumber's font-derived counts as the reference — but the extraction stage is ultimately judged by "
     "the full PDF-to-JSON score, since that is what the markers are for.")
todo(["Decide and state how the evaluation set will grow: with 8 section.description items and 16 "
      "question.text items, one document moves those labels by tens of points.",
      "Annotator agreement, if measured, belongs here as the ceiling the scores should be read against."])

# ── 9 ────────────────────────────────────────────────────────────────────
h1("9. Results")
para("All ten documents, 75% containment, both paths. Runs on disk as of 28 August 2026.", size=10,
     italic=True, color=GREY)
h2("9.1 Models — pdfplumber extractor")
table(["Model", "Precision", "Recall", "Path A F1", "Path B F1"],
      [["gemma4:e4b", "0.961", "0.932", "0.946", "0.951"],
       ["llama3.3:70b", "0.736", "0.826", "0.779", "0.767"],
       ["qwen2.5:14b", "0.679", "0.833", "0.748", "0.750"],
       ["llama3.1:8b", "0.619", "0.750", "0.678", "0.670"]],
      widths=[1.4, 1.1, 1.1, 1.1, 1.1])
table(["Label (Path A F1)", "gemma4:e4b", "llama3.3:70b", "qwen2.5:14b", "llama3.1:8b", "gold items"],
      [["title", "1.000", "1.000", "1.000", "1.000", "10"],
       ["section.title", "1.000", "0.956", "0.887", "0.778", "43"],
       ["section.description", "0.889", "0.457", "0.583", "0.182", "8"],
       ["question.text", "0.828", "0.706", "0.421", "0.065", "16"],
       ["answer.text", "0.935", "0.713", "0.730", "0.738", "55"]],
      widths=[1.5, 1.0, 1.05, 1.05, 1.05, 0.8], size=8.5)
para("question.text is the label that separates the models: gemma4:e4b finds 12 of the 16 real questions, "
     "llama3.3:70b 12, qwen2.5:14b 8, llama3.1:8b 1 (it sends 12 to section.title, identically across every "
     "prompt tried). section.description is every model's weakest label and has only 8 gold items.",
     size=10, color=GREY)

h2("9.2 Extractors — gemma4:e4b")
table(["Extractor", "Precision", "Recall", "Path A F1", "Path B F1"],
      [["pdfplumber", "0.961", "0.932", "0.946", "0.951"],
       ["Docling (native cells)", "0.975", "0.879", "0.924", "0.910"],
       ["LightOnOCR-2-1B", "0.823", "0.811", "0.817", "0.827"],
       ["Docling, Markdown export (superseded)", "0.832", "0.712", "0.767", "0.757"]],
      widths=[2.2, 1.0, 1.0, 1.0, 1.0])
table(["Sample", "pdfplumber", "Docling", "LightOnOCR", "What decides it"],
      [["1", "1.000", "1.000", "1.000", "NIH template — structure is in the words"],
       ["2", "0.818", "0.870", "0.533", "41 italic instruction paragraphs, inline bold labels"],
       ["3", "1.000", "1.000", "0.963", "headings by size"],
       ["4", "1.000", "1.000", "1.000", "two blocks"],
       ["5", "0.786", "0.923", "0.538", "inline bold labels"],
       ["6", "1.000", "0.154", "1.000", "five drawn-underline headings"],
       ["7", "1.000", "1.000", "1.000", "two blocks"],
       ["8", "1.000", "1.000", "0.533", "bold-italic headings"],
       ["9", "1.000", "1.000", "1.000", "headings by size"],
       ["10", "1.000", "1.000", "1.000", "italic headings, underlined title"]],
      widths=[0.6, 0.95, 0.85, 0.95, 3.0], size=8.5)
para("The markers are the whole game. pdfplumber and Docling read the same fonts and score within two "
     "points; the difference is one document (sample 6) whose headings are drawn underlines that Docling's "
     "data does not contain. Docling is ahead where inline labels and italic instructions matter (samples 2 "
     "and 5) and finds 13 of the 16 real questions to pdfplumber's 12. LightOnOCR, which guesses the markers "
     "from an image, is 13 points back.", size=10, color=GREY)
para("Recommendation: pdfplumber + gemma4:e4b — the highest score, no structural blind spot on this corpus, "
     "no layout model, sub-second extraction. Docling is a close second whose value is what pdfplumber lacks: "
     "block boundaries, a page and position for every block, and a confidence per block. LightOnOCR is the "
     "only one of the three that can read a scanned page.", size=10.5, bold=True)
para("Known limitation of the recommended configuration: bold labels sharing a line with their answer are "
     "sometimes kept fused into the answer (every question.text it misses is this pattern); the "
     "answer-versus-description boundary is a judgement the annotation draws and the text does not mark; "
     "and all of it rests on a text layer, ten documents and a prompt that two of them helped write.",
     size=10, italic=True, color=RUST)
todo(["The extractor table spans two prompt versions (section 7). Re-run and replace before publication.",
      "One test in the suite fails (sample 3: the annotation rules add a question stage 3 does not have). "
      "Resolve or explain before quoting Path B figures as final."])

# ── 10 ───────────────────────────────────────────────────────────────────
h1("10. Open Items — checklist")
table(["#", "Item", "Section"],
      [["1", "Write the literature review and add the numbered references", "2"],
       ["2", "Confirm or reword the RDA DMP Common Standard output claim", "1"],
       ["3", "Annotators, guidelines, agreement; funder and source per sample", "4"],
       ["4", "Re-run pdfplumber and LightOnOCR under the current prompt; decide the doubled dash", "7, 9"],
       ["5", "Resolve the sample-3 rules test failure", "9"],
       ["6", "Replace approximate VRAM figures with measured ones", "6"],
       ["7", "Decide whether block-by-block prompting will be tested, and whether the prompt examples from "
             "samples 1–2 stay", "7"],
       ["8", "Plan for growing the evaluation set beyond ten documents", "8"]],
      widths=[0.4, 5.1, 0.9], size=9)

out = Path("Report-doc/design_document.docx")
out.parent.mkdir(parents=True, exist_ok=True)
doc.save(out)
print(f"saved -> {out}  ({out.stat().st_size / 1024:.0f} KB)")
