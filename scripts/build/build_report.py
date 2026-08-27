"""Build Report-doc/project_report.docx as a clean, results-free design document.

All measured results have been removed: the prompt fix, line merging and the
move to a 10-sample evaluation set each invalidate the previous figures, and
the raw outputs have been deleted. What remains describes the pipeline as it
now stands, ready to be populated by a fresh run.
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH as _ALIGN
from docx.shared import Inches, Pt, RGBColor

TITLE_BLUE = RGBColor(0x1E, 0x40, 0x6E)
SUB_BLUE   = RGBColor(0x22, 0x63, 0x99)
H2_BLUE    = RGBColor(0x4F, 0x81, 0xBD)
GREY       = RGBColor(0x59, 0x59, 0x59)
RUST       = RGBColor(0x99, 0x3A, 0x2B)

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(10.5)
for s in doc.sections:
    s.left_margin = s.right_margin = Inches(0.98)
    s.top_margin = s.bottom_margin = Inches(0.87)


def shade(cell, hexcolor):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def setc(cell, text, *, bold=False, size=9, color=None, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color


def table(headers, rows, widths=None, size=9):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]
        setc(c, hd, bold=True, size=size, color=RGBColor(0xFF, 0xFF, 0xFF), center=i > 0)
        shade(c, "1E406E")
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, v in enumerate(row):
            setc(cells[ci], str(v), bold=(ci == 0), size=size, center=ci > 0)
            if ri % 2 == 1:
                shade(cells[ci], "F2F5FA")
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = TITLE_BLUE
    return p


def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = H2_BLUE
    return p


def para(text, *, size=10.5, bold=False, italic=False, color=None, after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return p


def mono(lines, size=8.5):
    for ln in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(ln)
        r.font.name = "Consolas"
        r.font.size = Pt(size)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def bullet(text, prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    if prefix:
        r = p.add_run(prefix)
        r.bold = True
        r.font.size = Pt(10.5)
    r = p.add_run(text)
    r.font.size = Pt(10.5)


# ── Header ───────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run("DMPbridge")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = TITLE_BLUE

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(10)
r = p.add_run("Automated Labeling Pipeline for Data Management Plan Documents")
r.font.size = Pt(13)
r.font.color.rgb = SUB_BLUE

for line in [
    "Date: 19 August 2026",
    "Dataset: 10 manually labeled DMP documents (23 pages) · all 10 scored",
    "Models: Gemma 4 e4b · Qwen 2.5 14B · Llama 3.3 70B · Llama 3.1 8B (via Ollama, free/local)",
    "Extractors: pdfplumber (default) · LightOnOCR-2-1B (alternative, gemma4:e4b only — see 6.1a)",
    "Experiments: 4 of 4 complete under the current pipeline",
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(line)
    r.bold = True
    r.font.size = Pt(11)

para("Results in section 6 reflect every fix through 19 August 2026 (section 6.1a), all four models "
     "complete under pdfplumber's whole-document pipeline (section 3.3), scored at the project's 75% "
     "partial-credit containment default. Section 4.1's inference-parameters table (format: OUTPUT_SCHEMA, "
     "and its 'nine configurations' framing) describes the pre-rewrite, id-based design and has not been "
     "updated to match — pdfplumber no longer uses OUTPUT_SCHEMA at all; treat that table as historical.",
     size=9.5, italic=True, color=RUST, after=10)

# ── 1 ────────────────────────────────────────────────────────────────────
h1("1. Project Overview")
para("DMPbridge is a Python pipeline for automatically labeling text blocks in Data Management Plan (DMP) PDF "
     "documents. It extracts structured text blocks from PDFs and uses large language models to classify each "
     "block with one of five structural labels, enabling downstream tasks such as DMP comparison, "
     "funder-specific analysis, and automated compliance checking.")
para("The pipeline has two interchangeable components. The extraction backend converts a PDF into text blocks — "
     "three are implemented. The labeling strategy sends those blocks to a model for classification; one "
     "strategy, whole-document, is currently implemented.")
para("All experiments run locally via Ollama — no cloud API cost, and no document leaves the machine.")

# ── 2 ────────────────────────────────────────────────────────────────────
h1("2. Label Schema")
table(["Label", "Description", "Written by", "Key signal"],
      [["title", "Document main title", "Researcher", "Largest font, first block"],
       ["section.title", "Section or sub-section heading", "Funder/researcher", "Numbered, bold, short"],
       ["section.description", "Funder template text / instructions", "Funder", "Uses 'should', 'must'"],
       ["question.text", "Sub-question or sub-topic prompt", "Funder", "Letter prefix (A., B.)"],
       ["answer.text", "Researcher's written response", "Researcher", "Narrative, first-person"]],
      widths=[1.35, 2.15, 1.25, 1.75])
para("Key challenge: question.text and section.description are the hardest pair — both are funder-written; the "
     "boundary is structural, not lexical.", size=10, italic=True, color=GREY)

# ── 3 ────────────────────────────────────────────────────────────────────
h1("3. Architecture")
h2("3.1 Module Layout")
mono([
    "dmpbridge/",
    "├── core/",
    "│   ├── pipeline.py            process_pdf() — top-level orchestrator",
    "│   ├── converter.py           to_structured() — flat blocks → DMP Tool schema",
    "│   ├── paths.py               the four-stage output layout",
    "│   └── config.py              defaults",
    "├── extractors/                pluggable PDF → blocks layer",
    "│   ├── base.py                BaseExtractor protocol",
    "│   ├── pdfplumber_extractor.py  — the only extractor implemented",
    "│   └── __init__.py            get_extractor() factory",
    "├── preprocess/",
    "│   ├── pdfplumber_reader.py   extract_text_for_llm() — whole doc, **bold**/_italic_ markers",
    "│   └── page_images.py         page PNGs with bbox overlays",
    "├── strategies/",
    "│   ├── wholedoc.py            WholeDocStrategy — the only strategy",
    "│   └── __init__.py            get_strategy() factory",
    "├── models/",
    "│   ├── ollama.py              OllamaModel — local models via Ollama REST",
    "│   └── __init__.py            get_model() factory",
    "├── parsers/json_parser.py     parse_llm_json()",
    "├── prompts/",
    "│   └── constants.py           LABELS + pdfplumber's own SYSTEM_PROMPT/schema",
    "├── evaluation/",
    "│   ├── evaluate.py            Path A — scoring vs old annotation",
    "│   ├── annotation_rules.py    Path B — rule conversion + scoring vs new",
    "│   └── experiment.py          ExperimentRunner CLI",
    "└── cli/",
    "    ├── main.py                dmpbridge",
    "    └── wholedoc_cmd.py        dmpbridge-wholedoc",
])

h2("3.2 Data Flow")
para("The pipeline runs in four stages, each writing its own directory so any stage can be inspected or "
     "re-run without disturbing the others. Filenames are identical at every stage, so one document can be "
     "followed from folder to folder by opening the same sampleN.json.")

_img = Path("Report-doc/pipeline_diagram.png")
if _img.exists():
    _p = doc.add_paragraph()
    _p.alignment = _ALIGN.CENTER
    _p.paragraph_format.space_before = Pt(6)
    _p.paragraph_format.space_after = Pt(4)
    _p.add_run().add_picture(str(_img), width=Inches(6.1))
    _cap = doc.add_paragraph()
    _cap.alignment = _ALIGN.CENTER
    _cr = _cap.add_run("Figure 1 — The four pipeline stages and the two evaluation paths. Extraction is "
                       "cached separately from labeling because it does not depend on the model.")
    _cr.font.size = Pt(8.5)
    _cr.italic = True
    _cr.font.color.rgb = GREY
    _cap.paragraph_format.space_after = Pt(10)

para("Two design points are worth drawing out of the diagram.")
bullet("Extraction does not depend on which model labels the blocks, so stage 1 is keyed by extractor and "
       "cached. Labeling a corpus with three models costs one extraction rather than three.",
       prefix="Stage 1 is shared. ")
bullet("Path A scores stage 3 against the original annotation; Path B scores stage 4, after the Rules.xlsx "
       "conversion, against the newer one. Both stages are always written, so the converted and unconverted "
       "forms can be compared directly. Earlier revisions overwrote stage 3 in place, and only one of the two "
       "could exist at a time.",
       prefix="The paths diverge at stage 3. ")

h2("3.3 Block Schema")
para("Extraction stage 1 output — one entry, the whole document:", size=9.5, italic=True, color=GREY)
table(["Field", "Type", "Description"],
      [["text", "str", "Whole document, with **bold**/_italic_ visual-signal markers embedded"]],
      widths=[1.4, 1.2, 3.9])
para("Classification stage 2 output — one entry per classified item:", size=9.5, italic=True, color=GREY)
table(["Field", "Type", "Description"],
      [["text", "str", "Verbatim source text, markers stripped"],
       ["label", "str", "One of the five labels"]],
      widths=[1.4, 1.2, 3.9])

# ── 4 ────────────────────────────────────────────────────────────────────
h1("4. Labeling Strategy and Extraction Backends")
h2("4.1 Labeling strategy — whole-document")
para("All blocks are sent in a single model call with full document context, using a grammar-constrained JSON "
     "schema so the response is always valid. This is the only strategy in the codebase; the Batch strategy "
     "described in the July 2026 revision has been removed.")
para("This variable is therefore held constant across every experiment rather than chosen after comparison. No "
     "evidence has been collected that whole-document is the best approach. It also means document length "
     "scales directly with context usage — the corpus here is short, the longest document being 5 pages, so "
     "that limit has never been approached and remains untested.", size=10, color=GREY)

para("Inference parameters — identical across all nine configurations:", size=10, after=4)
table(["Parameter", "Value", "Rationale"],
      [["format", "OUTPUT_SCHEMA", "Ollama grammar-constrains output to the JSON schema, so malformed "
                                   "responses are impossible and no repair step is needed"],
       ["temperature", "0.0", "Deterministic decoding"],
       ["num_ctx", "32768", "Comfortably holds a full DMP for whole-document inference"],
       ["keep_alive", "-1", "Model stays resident in VRAM between documents, avoiding reload per sample"],
       ["timeout", "3600 s", "Accommodates the 70B model on long documents"]],
      widths=[1.15, 1.15, 4.1], size=8.5)
para("Decoding is deterministic and the output grammar is enforced. This was verified on 6 August by running "
     "one configuration three times: every count was identical and F1 did not move at all. The measured "
     "run-to-run variation is +/-0.002 F1 — a single text block. Differences above roughly 0.005 F1 between "
     "configurations are therefore real signal, not noise. See section 4.2.",
     size=10, italic=True, color=GREY)

h2("4.2 Label definitions in the prompt")
para("The prompt's wording is not incidental — one line of it was measurably wrong until 6 August.")
para("Both models' single largest error was labelling lettered sub-items ('B. Scientific data that will be "
     "preserved…') as section.title rather than question.text: 12 occurrences for llama3.1:8b, 10 for "
     "gemma4:e4b. The cause was that the prompt described section.title as 'often starts with a letter prefix "
     "(A., B., C.)', which is the opposite of what the reference data does:")
table(["Heading pattern", "section.title", "question.text"],
      [["Letter prefix (A., B., C.)", "0", "8"],
       ["Number prefix (1., 2.)", "13", "0"],
       ["Starts 'Element N:'", "6", "0"]],
      widths=[2.6, 1.9, 1.9])
para("Letter prefixes are used exclusively for questions. The description also contradicted its own worked "
     "example, which was 'A. Types and amount of scientific data…' — and the models followed the description.")
para("section.title now states that it does not include lettered sub-items; question.text states that it "
     "frequently is one. The effect divided sharply by model:")
table(["Model", "Path A F1 before", "after", "question.text F1", "Lettered items mislabelled"],
      [["gemma4:e4b", "67.6%", "74.3%", "16.2% -> 41.9%", "5/8 -> 0/8"],
       ["llama3.1:8b", "61.7%", "61.9%", "4.1% -> 4.2%", "7/8 -> 7/8"]],
      widths=[1.2, 1.35, 0.85, 1.55, 1.45], size=8.5)
para("The correction works completely on gemma4:e4b and not at all on llama3.1:8b, which still mislabels "
     "seven of eight lettered items. The instruction is present; the smaller model does not act on it. "
     "Negative constraints appear to require more capability than positive ones — a consideration when "
     "writing prompts for models at this scale.", size=10, color=GREY)
para("This also explains an earlier unknown: removing a duplicated prompt on 5 August cost roughly 16 points "
     "on pdfplumber for no visible reason. The deleted copy carried exactly this guidance.",
     size=10, color=GREY)

para("The noise floor, measured", bold=True, size=11)
para("On 6 August the same configuration (llama3.1:8b, pdfplumber) was run three times with an identical "
     "prompt. Every count was identical — TP 96, FP 89, FN 36 — and every per-class F1 agreed to three "
     "decimals. Over four runs the largest variation observed anywhere was a single block.")
para("The run-to-run noise floor is +/-0.002 F1. Earlier drafts of this report assumed roughly 3 points and "
     "used that figure to argue the top two models could not be separated; that assumption was never measured "
     "and is now withdrawn. Differences above roughly 0.005 F1 are real.")

para("Prompt variants, and a reversal", bold=True, size=11)
para("With a measurable floor, four wordings of the section.title definition were compared on llama3.1:8b. "
     "All four differ by more than the noise floor, so the ordering is meaningful:")
table(["section.title definition", "Path A F1"],
      [["'Often starts with a letter prefix (A., B., C.) or a bold named phrase'", "62.5%"],
       ["Original wording", "61.9%"],
       ["As the first row, but with blank lines split differently", "60.6%"],
       ["Explicitly excludes lettered sub-items (the 4.2 correction)", "59.1%"]],
      widths=[4.6, 1.4], size=8.5)
para("This reverses the conclusion above for this model. The correction described in 4.2 is the worst of the "
     "four on llama3.1:8b, and restoring the letter-prefix claim is the best. The evidence about the reference "
     "data still stands — letter prefixes really are used only for questions — but the 8B model does not "
     "respond to that instruction as intended in either direction.")
para("Rows one and three have identical wording and differ only in the placement of three blank lines, yet "
     "differ by 1.9 points of F1 and 13 false positives. Both layouts were re-run and each reproduced to "
     "within one block. The prompt must therefore be treated as whitespace-significant: an editor that strips "
     "or adds a blank line on save can move the score more than most rewording does.")
para("Prompt changes are also model-specific and do not transfer. The best llama3.1:8b variant above was run "
     "on gemma4:e4b and cost it 4.0 points (74.3% to 70.3%) — not on question.text, which was unchanged, but "
     "on title, which fell from 90.0% to 77.8% as bold document titles were drawn into section.title. Any "
     "prompt change requires re-running every model.")
para("One result has survived every variant: across eight runs and four prompts, llama3.1:8b sends 12 of the "
     "16 real questions to section.title, 3 to nothing, and labels 1 correctly — identical every time. "
     "gemma4:e4b and llama3.3:70b each label 9 of 16 correctly using the same prompt. This is a capability "
     "limit rather than a prompt defect.", size=10, color=GREY)

h2("4.3 Extraction backends")
table(["Backend", "Method", "Bbox / font", "Blocks/doc", "Requires"],
      [["pdfplumber", "Whole-document text, extraction+labeling fused", "No", "1", "nothing (bundled)"],
       ["LightOnOCR-2-1B", "Vision-LLM OCR, same marker convention as pdfplumber", "No", "1",
        "CUDA GPU, torch/transformers (dmpbridge[lighton])"],
       ["Docling", "Layout model + native page cells: pdfplumber's font rules on Docling's words, "
                   "hyperlinks as underline, headings from layout", "No", "1",
        "docling (dmpbridge[docling]), CPU-capable"]],
      widths=[1.25, 2.0, 1.05, 1.05, 1.65])
para("Docling was tried, removed, and re-added 24 August 2026 in whole-document form. Its first "
     "version worked from Docling's Markdown export, which carries no bold/italic/underline, and scored "
     "0.767 Path A / 0.757 Path B with gemma4:e4b. On 27 August it was rebuilt to read Docling's native "
     "page cells (font names, sizes, hyperlink rectangles) with the same rules pdfplumber applies to its "
     "characters: markers now match pdfplumber's on 8 of 10 documents and it scores 0.924 / 0.910 — "
     "ahead of pdfplumber on samples 2 and 5, level on seven, behind only on sample 6, whose drawn "
     "underlines Docling has no shape data for. LightOnOCR was tried, "
     "removed, then re-added 19 August 2026 as a working, non-default alternative — see 6.1a for the "
     "current pdfplumber-vs-LightOnOCR comparison. pdfplumber remains the default given both its accuracy "
     "and speed advantage.",
     size=10, italic=True, color=GREY)

# ── 5 ────────────────────────────────────────────────────────────────────
h1("5. Experiment Registry")
para("Five configurations: four models against pdfplumber (the default), plus gemma4:e4b against "
     "LightOnOCR-2-1B (the alternative extractor, section 4.3/6.1a), whole-document strategy, scored on "
     "all ten samples through both evaluation paths.")
rows=[]; eid=1
DONE={("llama3.3:70b","pdfplumber"),("gemma4:e4b","pdfplumber"),("llama3.1:8b","pdfplumber"),
      ("qwen2.5:14b","pdfplumber"),("gemma4:e4b","lightonocr")}
CONFIGS = [("llama3.3:70b","llama3.3-70b","pdfplumber"),
           ("gemma4:e4b","gemma4-e4b","pdfplumber"),
           ("llama3.1:8b","llama3.1-8b","pdfplumber"),
           ("qwen2.5:14b","qwen2.5-14b","pdfplumber"),
           ("gemma4:e4b","gemma4-e4b","lightonocr")]
for mo, slug, ex in CONFIGS:
    rows.append([f"E{eid:02d}", mo, ex, f"{slug}_{ex}_whole_doc",
                 "complete" if (mo,ex) in DONE else "pending"])
    eid+=1
table(["ID","Model","Extractor","Output tag","Status"], rows, widths=[0.5,1.3,1.1,2.6,0.85])
para("All five configurations are complete under the current pipeline as of 19 August 2026.",
     size=9.5, italic=True, color=GREY)

# ── 6 ────────────────────────────────────────────────────────────────────
h1("6. Results")
para("All ten samples, pdfplumber's whole-document visual-signal pipeline (section 4), current annotation "
     "rules, all fixes through 19 August 2026 (section 6.1a). Produced 19 August 2026 — supersedes every "
     "earlier results table in this report, including the 18 August figures previously here.",
     size=10, italic=True, color=RUST)
para("Scored at the project default containment threshold — 75% (CONTAINMENT_THRESHOLD = 0.75 in "
     "evaluate.py) — a predicted item matches once at least 75% of its words appear in the reference item, "
     "partial credit allowed. The project ran at an exact-match 100% threshold from 12 August to 19 August "
     "2026; any figure elsewhere describing a \"100%\" or \"exact containment\" default predates this and is "
     "not directly comparable to what follows.", size=9.5, italic=True, color=GREY)

h2("6.1 Headline")
table(["Model","Path A F1","Path B F1"],
      [["gemma4:e4b","94.6%","94.8%"],
       ["llama3.3:70b","77.9%","76.4%"],
       ["qwen2.5:14b","74.8%","74.7%"],
       ["llama3.1:8b","67.8%","66.7%"]],
      widths=[1.35,1.35,1.35])
para("gemma4:e4b remains the clear best model, now scoring above 94% on both paths. Ranking below it has "
     "reordered since 18 August: llama3.3:70b is now a clear second (77.9%/76.4%, helped substantially by "
     "the fixes in 6.1a), qwen2.5:14b third, llama3.1:8b last — llama3.3:70b's earlier apparent tie with "
     "llama3.1:8b (18 August) has not held up under corrected extraction and a broader set of fixes.",
     size=10, color=GREY)
para("Path A and Path B are now close on every model — within 1.5 percentage points, either direction — "
     "a marked change from 18 August, where Path B trailed Path A by 2.7-3.8pp across the board. See 6.3.",
     size=10, italic=True, color=RUST)

h2("6.1a What changed since 18 August")
para("Five separate, real fixes, each found via a direct user-reported discrepancy and root-caused before "
     "being applied — not guessed at. Full detail: Report-doc/worklog/2026-08-19.md (not published; kept "
     "locally only).", size=10, color=GREY)
bullet("a PDF-authoring artifact (a colon rendered in a different font than its own sentence, with no real "
       "bold styling) was being flagged as emphasized text, corrupting sample1's structure. Fixed by "
       "requiring a punctuation-only word's substituted font to actually be bold-named, not just different.",
       prefix="Font-mismatch false positive — ")
bullet("sample6's five section headings are underlined in the source PDF, which font-based bold/italic "
       "detection cannot see at all (underline is a drawn shape, not a font attribute). Added detection "
       "against the page's drawn rectangles and a new ++underline++ marker.",
       prefix="Underline detection — ")
bullet("the underline fix's first version also marked hyperlinks and incidental emphasis as \"underlined\" "
       "across five other samples, destabilizing gemma's whole-document generation on unrelated content "
       "(sample2 dropped 0.769 to 0.522). Excluding URL-shaped text from underline detection not only cut "
       "the noise but fully recovered sample2 to its original score.",
       prefix="URL exclusion — ")
bullet("a character-doubling rendering artifact (\"aa\" instead of \"a\") was too short to trigger the "
       "existing char-dedup fix, cascading into two separate scoring errors on sample3. Fixed with a "
       "context-aware check — a short word is only collapsed when its line's longer words show the same "
       "doubling pattern, which correctly leaves sample2's genuine \"XX\" placeholder text untouched.",
       prefix="Character-dedup gap — ")
bullet("CONTAINMENT_THRESHOLD reverted from the exact-match 100% (in effect 12-19 August) back to the "
       "original 75% partial-credit default, by direct decision.",
       prefix="Threshold — ")
para("A vision-based OCR extractor (LightOnOCR-2-1B) was also fully integrated as an alternative to "
     "pdfplumber and tested against gemma4:e4b across all 10 samples: it independently confirms sample6's "
     "underline fix (also scores a perfect 1.000 there, via a different mechanism), but scores lower "
     "overall (81.3% vs. pdfplumber's 90.7% pooled F1, both at the time compared) and runs roughly 14x "
     "slower. Kept as a working, non-default extractor rather than adopted — "
     "notebooks/comparison-gemma-pdfplumber-vs-lightonocr.ipynb has the full comparison.",
     size=10, italic=True, color=GREY)

h2("6.2 Per-label F1")
table(["Label","gemma4:e4b (A/B)","qwen2.5:14b (A/B)","llama3.3:70b (A/B)","llama3.1:8b (A/B)","Gold (A/B)"],
      [["title","100/100%","100/100%","100/100%","100/100%","10/10"],
       ["section.title","100/100%","88.7/88.0%","95.6/94.6%","77.8/77.5%","43/45"],
       ["section.description","88.9/88.9%","58.3/58.3%","45.7/45.7%","18.2/18.2%","8/8"],
       ["question.text","82.8/91.7%","42.1/67.2%","70.6/70.6%","6.5/48.8%","16/56"],
       ["answer.text","93.5/93.5%","73.0/69.6%","71.3/71.3%","73.8/73.8%","55/55"]],
      widths=[1.4,1.15,1.15,1.15,1.15,0.85], size=8.5)
para("question.text recovering sharply from Path A to Path B for every model except llama3.3:70b (flat) "
     "and qwen2.5:14b's answer.text dipping slightly on Path B are the annotation rules working as "
     "designed — Path A scores against the original annotation, which leaves many questions blank; Path B "
     "fills them first. section.description remains every model's weakest label by a wide margin "
     "(18.2-88.9%, base of only 8 gold items) — the smallest, noisiest label in the schema, unchanged "
     "since 18 August.", size=10, color=GREY)

h2("6.3 Path A versus Path B")
para("Path A and Path B are now close on every model: llama3.1:8b -1.1pp, gemma4:e4b +0.2pp, "
     "llama3.3:70b -1.5pp, qwen2.5:14b -0.1pp — a real change from 18 August, where Path B trailed Path A "
     "by 2.7-3.8pp on every model with no exception. The gap has not fully closed and is not expected to "
     "(Path A and Path B are scored against two separately-annotated reference sets, not the same answer "
     "key scored twice — see 7.3), but it is now small and mixed in direction rather than a one-sided, "
     "every-model pattern.", size=10, color=GREY)

h2("6.4 Runtime — why the new pipeline runs slower")
para("pdfplumber's extraction and labeling were rewritten (section 3.3) to match a whole-document, "
     "visual-signal design: the model reads the entire document as one text string with **bold**/_italic_ "
     "markers standing in for visual emphasis, and returns a flat [{\"text\", \"label\"}] array directly, "
     "rather than being handed a pre-segmented list of blocks and asked to classify each by id.")
table(["Model", "Old pipeline, 10 docs", "New pipeline, 10 docs", "Change"],
      [["llama3.1:8b", "74 s", "149 s", "2.0×"],
       ["gemma4:e4b", "74 s", "157 s", "2.1×"],
       ["qwen2.5:14b", "n/a — new to this pipeline", "235 s", "—"],
       ["llama3.3:70b", "~20 min (documented)", "~40+ min", "~2×"]],
      widths=[1.3, 1.6, 1.8, 1.0])
para("The cause is architectural, not incidental. The old pipeline's output was terse — one "
     "{id, label, confidence} triple per block, with the model never regenerating any of the document's "
     "actual text. The new pipeline's output includes the classified text itself for every item, so the "
     "model must reproduce the full source content as output tokens, not just short labels. LLM generation "
     "speed is bottlenecked by output token count far more than by input size, so this is a real, roughly "
     "twofold increase in generation work per document, consistent across every model measured — not GPU "
     "contention, and not specific to one document.", size=10, color=GREY)
para("One document-level effect compounded this on llama3.3:70b specifically during this run: sample2's "
     "extracted text is 14,898 characters, more than double every other sample in the set. On a 70B model, "
     "where per-token latency is already far higher than the smaller models, an unusually long document adds "
     "directly to an already-slower-per-document baseline.", size=10, italic=True, color=GREY)
para("Practical consequence: the ~112 s/sample, ~20-minute figure for llama3.3:70b recorded in the "
     "project's own operating notes was measured under the old, pre-rewrite pipeline and is now stale for "
     "the same architectural reason the old score table was — expect roughly double that under the current "
     "pipeline going forward.", size=10, italic=True, color=RUST)
para("A separate, operational cause of slow runs was found and fixed 19 August: the Ollama server process "
     "itself had been running continuously for 7 days (since 12 August) across every model swap in that "
     "window, and had silently degraded — one sample took 226s against an established ~15s baseline, with "
     "no architectural explanation. Restarting Ollama cleanly (with the environment flags this project's "
     "own operating notes already specify) restored the baseline immediately. Not a pipeline defect — a "
     "reminder that a long-lived Ollama process should be restarted periodically, not just when a model "
     "swap seems to be failing.", size=10, italic=True, color=GREY)

# ── 7 ────────────────────────────────────────────────────────────────────
h1("7. Evaluation Methodology")
h2("7.1 Matching")
para("Predictions and ground truth are both reduced to a flat list of (text, label) pairs from the structured "
     "JSON. Each gold item greedily claims the best unclaimed predicted item by token containment (≥ 0.75). "
     "A predicted item claimed by no gold item is recorded as a false positive. One matching routine feeds the "
     "confusion matrix, the error listing and the missed-item report, so those can never disagree.")

h2("7.2 Metrics")
table(["Metric", "Definition", "Penalises"],
      [["Precision", "TP / all predicted items", "Spurious / over-generated blocks"],
       ["Recall", "TP / all gold items", "Missed gold items"],
       ["Micro-F1", "Harmonic mean of the two", "Both — used for all ranking"]],
      widths=[1.15, 2.35, 2.9])
para("The July 2026 revision reported block-level and gold-based accuracy. Micro-averaged precision/recall/F1 "
     "replaced those because accuracy does not penalise over-generation. Figures from that revision are not "
     "comparable with anything reported since.", size=10, italic=True, color=GREY)

h2("7.3 Two evaluation paths")
para("The manual annotation standard changed partway through the project, so every configuration is scored "
     "twice against two independent reference sets:")
bullet("scores the structured JSON directly against data/input/ground_truth_old_version/.", prefix="Path A ")
bullet("applies apply_new_annotation_rules() — filling blank question text from the section title, or the "
       "document title as fallback — then scores the result against data/input/ground_truth_new_version/.",
       prefix="Path B ")
para("Both paths share the same matching and scoring engine, so any difference reflects the reference data and "
     "the conversion rule, not the measurement.")
para("The conversion is specified by data/input/Rules.xlsx — a 16-row truth table over whether each of "
     "title, section.title, section.description and question.text is empty (E) or non-empty (N). Exactly one "
     "row matches any question. The library transcribes the sheet directly rather than interpreting it.")
para("The sixteen rows express one rule in two halves:", size=10, after=4)
bullet("leave it alone — every even row.", prefix="question.text already has text ")
bullet("fill it from the first available source: section.title, then section.description, then the document "
       "title.", prefix="question.text is empty — ")
para("question.text is the only field the rule writes; neither section.title nor the document title is "
     "modified. The rule reproduces the new-version reference annotation on all ten documents.")
para("Column order is load-bearing. On 6 August the sheet's columns were reordered — question.text moved from "
     "column 2 to column 4 — while the action text stayed identical. Nothing looked broken, but six rows were "
     "silently remapped and agreement with the reference files fell from 10/10 to 2/10. The library now "
     "records the expected order in RULE_FIELDS, and both the test suite and "
     "notebooks/annotation_conversion_test.ipynb assert the sheet's header against it.",
     size=10, color=GREY)
para("The sub-question merge described in earlier revisions of this report — several short sub-questions "
     "collapsed into one with answers concatenated — is no longer an open problem. Under the current table "
     "sample 5 matches exactly and no merging is required.", size=10, color=GREY)

h2("7.4 Evaluation set (changed 5 August)")
para("Samples 1 and 2 were originally held out because they were used during prompt development. From "
     "5 August all ten documents are scored; the notebooks set EXCLUDE = [].")
para("The reason is that those two documents are the NIH-style plans with explicit lettered sub-questions, and "
     "they hold most of the question.text ground truth in the corpus:")
table(["Evaluation set", "Gold items (old gt)", "…question.text", "Gold items (new gt)", "…question.text"],
      [["Samples 3–10 (to 4 Aug)", "90", "6", "81", "2"],
       ["Samples 1–10 (from 5 Aug)", "143", "22", "134", "18"]],
      widths=[1.75, 1.25, 1.3, 1.2, 1.3], size=8.5)
para("On the 8-sample set question.text was effectively unmeasurable — with 2 gold items in the new annotation "
     "a single match moved the score by tens of points. With all ten documents there are 18, which is still "
     "small but no longer meaningless.")
para("The trade-off is explicit: samples 1 and 2 informed the prompt, so scores that include them are "
     "optimistic and are not a clean estimate of performance on unseen documents. The judgement is that a "
     "measurable question.text is worth more than a held-out pair at this stage, given that the label is one "
     "of the five the pipeline exists to produce.", size=10, color=GREY)

# ── 8 ────────────────────────────────────────────────────────────────────
h1("8. Reproduction")

h2("8.1 Python API")
para("process_pdf() runs extraction, labeling, structuring and optional rule conversion in one call:",
     size=10, after=4)
mono([
    "import dmpbridge",
    "",
    "blocks = dmpbridge.process_pdf(",
    '    "document.pdf",',
    '    model="gemma4:e4b",        # any Ollama tag',
    '    extractor="pdfplumber",    # the only extractor implemented',
    "    apply_rules=True,          # apply Path B conversion to structured output",
    '    output="labeled.json",',
    '    structured_output="structured.json",',
    "    raw_dir=None,              # skip saving pre-label blocks",
    ")",
])
table(["Parameter", "Default", "Effect"],
      [["extractor", '"pdfplumber"', "Selects the extraction backend"],
       ["apply_rules", "False", "Applies apply_new_annotation_rules() before writing structured output"],
       ["strategy", "None", "Pass a prebuilt Strategy; other model kwargs are then ignored"],
       ["raw_dir", '"data/output/extracted"', "Where pre-label blocks are saved; None to skip"],
       ["images_dir", "None", "Per-page PNGs with bbox overlays; None to skip"]],
      widths=[1.25, 1.5, 3.65], size=8.5)
para("Line merging is set on the extractor rather than through process_pdf(). It defaults to on:",
     size=10, after=4)
mono([
    "from dmpbridge.extractors import get_extractor",
    "",
    'new = get_extractor("pdfplumber")                      # 25.1 blocks/doc (default)',
])

h2("8.2 Full run — all three configurations")
mono([
    "python -m venv venv && venv\\Scripts\\activate",
    "pip install -e .",
    "",
    "ollama pull llama3.3:70b",
    "ollama pull gemma4:e4b",
    "ollama pull llama3.1:8b",
    "",
    "# start the server — on multi-GPU hosts see section 9 first",
    "CUDA_VISIBLE_DEVICES=0,1 OLLAMA_VULKAN=0 OLLAMA_SCHED_SPREAD=0 ollama serve",
    "",
    "# three runs, one per model",
    "for M in llama3.3:70b gemma4:e4b llama3.1:8b; do",
    "    dmpbridge-wholedoc --model $M --start 1 --end 10",
    "done",
    "",
    "# evaluate (all 10 samples — no --exclude)",
    "dmpbridge-evaluate      gemma4-e4b_pdfplumber_whole_doc   # Path A, reads stage 3",
    "dmpbridge-evaluate-new  gemma4-e4b_pdfplumber_whole_doc   # Path B, reads stage 4",
])
para("Existing output is skipped, so data/output/2_labeled/<tag>/ must be empty for a configuration to "
     "re-run. Stage 1 is cached separately and is reused unless --no-cache is passed. No API keys are "
     "required.", size=10, italic=True, color=GREY)
table(["Flag", "Effect"],
      [["--no-cache", "Re-extract even when stage 1 already holds the sample"],
       ["--no-rules", "Skip stage 4 — write the structured output but not the rule-converted version"],
       ["--start / --end", "Restrict to a range of samples"]],
      widths=[1.5, 4.9], size=8.5)
para("Extraction alone, with no model involved:", size=10, after=4)
mono([
    "python scripts/extract_pdfplumber.py",
    "python scripts/extract_pdfplumber.py --raw      # line-level, no merging",
    "python scripts/extract_pdfplumber.py --both     # both, with a comparison",
])

h2("8.3 Notebooks and tests")
table(["Artifact", "Purpose"],
      [["notebooks/01-run_pipeline.ipynb", "Run the pipeline on one PDF and inspect the structured output"],
       ["notebooks/1-model_comparison_pdfplumber.ipynb", "Full evaluation, pdfplumber"],
       ["notebooks/annotation_conversion_test.ipynb", "Derivation of the Path B rule"],
       ["tests/", "tests — converter, scoring engine, Path B rule, batch runner, per-sample invariants"]],
      widths=[2.7, 3.7], size=8.5)
mono(['pip install -e ".[dev]"', "pytest tests/"])

# ── 9 ────────────────────────────────────────────────────────────────────
h1("9. GPU Configuration")
para("This is the configuration the pipeline is run with. It is documented in full because the default "
     "behaviour on this hardware is unstable, and because the failure mode is silent rather than obvious.")

h2("9.1 Hardware")
table(["Component", "Value"],
      [["GPUs", "4 × NVIDIA GeForce RTX 3090, 24 GB each"],
       ["Driver", "591.86 (CUDA 13.1)"],
       ["Interconnect", "PCIe — no NVLink (consumer cards)"],
       ["Host", "Windows 11 Pro"]],
      widths=[1.6, 4.8])

h2("9.2 The problem")
para("Ollama detects both a CUDA and a Vulkan backend on this machine and, left to itself, selects Vulkan. Two "
     "consequences follow, and neither announces itself as an error:")
bullet("Vulkan multi-GPU inference is unstable here. The server died mid-run repeatedly, most often on the 70B "
       "model, surfacing to the client only as a dropped connection on port 11434.",
       prefix="Crashes. ")
bullet("Vulkan enumerates devices independently of CUDA_VISIBLE_DEVICES. Setting that variable alone correctly "
       "restricted CUDA to the cards at PCI 01:00 and 06:00, while Vulkan went on to claim 09:00 and 0a:00 — "
       "two different physical cards. All four GPUs showed memory in use despite the pin appearing to apply.",
       prefix="Device pinning silently ignored. ")
para("Disabling Vulkan resolved both, and roughly halved inference time on the 70B model.")

h2("9.3 Working procedure")
para("Step 1 — stop any running server. Ollama on Windows may already be running as a background app, and "
     "environment variables only apply to a freshly started process.", size=10, after=4)
mono(["taskkill /F /IM ollama.exe", 'taskkill /F /IM "ollama app.exe"'])
para("Step 2 — start the server with Vulkan disabled and CUDA pinned.", size=10, after=4)
para("Command Prompt:", size=9.5, bold=True, after=2)
mono(["set CUDA_VISIBLE_DEVICES=0,1", "set OLLAMA_VULKAN=0",
      "set OLLAMA_SCHED_SPREAD=0", "ollama serve"])
para("Bash / Git Bash:", size=9.5, bold=True, after=2)
mono(["CUDA_VISIBLE_DEVICES=0,1 OLLAMA_VULKAN=0 OLLAMA_SCHED_SPREAD=0 ollama serve"])
table(["Variable", "Purpose"],
      [["OLLAMA_VULKAN=0", "Forces the CUDA backend. The essential one — without it the other two have no "
                           "effect on the backend actually used."],
       ["CUDA_VISIBLE_DEVICES=0,1", "Restricts Ollama to two GPUs. 48 GB holds the 70B model (~42 GB); a "
                                    "single GPU suffices for the smaller models."],
       ["OLLAMA_SCHED_SPREAD=0", "Prevents spreading a model across more GPUs than it needs."]],
      widths=[1.9, 4.5])
para("Step 3 — leave that window open. The server runs in it; run the pipeline from a second window.", size=10)

h2("9.4 Verifying it worked")
para("Do not assume the settings applied — check the server log. Every 'inference compute' line must report "
     "library=CUDA, and no Vulkan device should be listed:", size=10, after=4)
mono([
    "# expected — CUDA only",
    'msg="inference compute" library=CUDA compute=8.6 name=CUDA0 ... pci_id=0000:01:00.0',
    'msg="inference compute" library=CUDA compute=8.6 name=CUDA1 ... pci_id=0000:06:00.0',
    'msg="selecting GPU backend for llama-server model" library=CUDA gpu_count=1',
    "",
    "# wrong — Vulkan was selected, crashes will follow",
    'msg="selecting GPU backend for llama-server model" library=Vulkan gpu_count=2',
])
para("Also check ollama ps once the model has loaded. It must report 100% GPU — anything less means layers "
     "have spilled to CPU and a 70B run will take hours rather than minutes. On 6 August a stale llama-server "
     "from an earlier session was still holding roughly 12 GB on each card, leaving only ~46 GB for a model "
     "needing ~54 GB; the run loaded at 15% CPU and had not finished one document in twelve minutes. After "
     "killing the stale process it loaded fully on GPU and averaged 132 s per document.",
     size=10, color=GREY)

h2("9.5 Allocation across components")
para("pdfplumber, the default extractor, is CPU-only, so there is no GPU contention to manage between "
     "extraction and labeling when using it. LightOnOCR-2-1B (re-added 19 August 2026, section 4.3) does "
     "need a GPU for its own inference (via torch/transformers, separate from Ollama's), and should be "
     "unloaded between uses the same way an idle Ollama model should be — see 9.6.")
table(["Component", "GPUs", "How it is set", "Notes"],
      [["Ollama (LLM)", "0, 1", "CUDA_VISIBLE_DEVICES on the server process",
        "70B needs both; smaller models fit on one"],
       ["pdfplumber", "—", "n/a", "CPU only"]],
      widths=[1.3, 0.7, 2.3, 2.1], size=8.5)

h2("9.6 Troubleshooting")
table(["Symptom", "Cause", "Fix"],
      [["ConnectionRefusedError on port 11434 mid-run",
        "Server crashed — check whether the log says library=Vulkan", "Restart with OLLAMA_VULKAN=0"],
       ["All 4 GPUs show memory despite pinning",
        "Vulkan backend in use; it ignores CUDA_VISIBLE_DEVICES", "OLLAMA_VULKAN=0"],
       ["'bind: Only one usage of each socket address'",
        "A server is already running", "Use it as-is, or taskkill first to apply new variables"],
       ["404 Not Found on /api/generate",
        "Model name wrong — not a connection problem", "Ollama tags use a colon: llama3.1:8b"],
       ["Out of memory loading 70B",
        "Only one GPU visible; the model needs ~54 GB", "CUDA_VISIBLE_DEVICES=0,1,2,3"],
       ["70B far slower than expected; ollama ps shows e.g. '15%/85% CPU/GPU'",
        "A stale llama-server from an earlier session still holds VRAM, leaving too little for the model",
        "taskkill /F /IM llama-server.exe, restart Ollama, confirm ollama ps shows 100% GPU"]],
      widths=[2.1, 2.2, 2.1], size=8.5)

# ── 10 ───────────────────────────────────────────────────────────────────
h1("10. Engineering Notes")

h2("10.1 Silent-failure defects")
para("Three defects produced silently wrong results rather than errors, and are recorded so they are not "
     "reintroduced.")
table(["Issue", "Cause", "Resolution"],
      [["LightOnOCR returned 'no text found' on every page",
        "Checkpoint declares a custom architecture but resolves to Mistral3, whose submodule names differ; "
        "the vision tower loaded as random weights with no error raised",
        "Remap weight keys on load; hard-fail if any weight is missing"],
       ["LightOnOCR emitted only '!!!!'",
        "Qwen3 text backbone is numerically unstable at float16", "Load in bfloat16"],
       ["Ollama crashed mid-run, worst with 70B",
        "Server selected the Vulkan backend over CUDA; Vulkan enumerates GPUs independently of "
        "CUDA_VISIBLE_DEVICES, so device pinning silently had no effect",
        "OLLAMA_VULKAN=0 plus CUDA pinning; also roughly twice as fast"]],
      widths=[1.6, 2.6, 2.2], size=8.5)

h2("10.2 Correctness fixes")
table(["Area", "Issue", "Resolution"],
      [["core/pipeline.py", "process_pdf() called extract_blocks() directly before delegating to the strategy, "
                            "so every PDF was parsed twice and the first parse ignored the requested extractor",
        "Removed; extraction delegated to the strategy alone"],
       ["evaluation/", "Two independent _snum() implementations; the one in evaluate.py used brittle string "
                       "replacement and failed on filenames not matching one fixed pattern",
        "Single regex implementation, imported by annotation_rules.py"],
       ["evaluation/", "annotation_rules.load_method() shadowed evaluate.load_method() — same name, different "
                       "ground truth and prediction directory",
        "Renamed to load_method_new(); callers, tests and notebooks updated"],
       ["evaluation/", "Path B rule cleared the document title after copying it into an empty question; the "
                       "reference files were revised on 5 Aug to retain it",
        "Title now left in place; rule again matches samples 1, 3, 4, 7 exactly"],
       ["experiments/", "A duplicate wholedoc.yaml covering all three models alongside the per-model YAMLs "
                        "produced repeated rows in benchmark output", "Duplicate removed"],
       ["models/ollama.py", "num_gpu=-1 forced every layer across all visible devices",
        "Removed; Ollama auto-selects (section 9)"],
       ["prompts/system.py", "section.title was described as 'often starts with a letter prefix (A., B., C.)' "
                             "— the opposite of the reference data, where letter prefixes mark questions",
        "Corrected; worth +6.7 F1 on gemma4:e4b (section 4.2)"],
       ["evaluation/", "Rules.xlsx column order changed without the transcription following, silently "
                       "remapping six rows and dropping agreement from 10/10 to 2/10",
        "RULE_FIELDS records the order; the test suite asserts the sheet header against it"],
       ["evaluation/", "Old ground-truth files renamed (sampleN_old_dmp.json -> sampleN_dmp_old.json), so "
                       "Path A matched zero files",
        "resolve_old_gt_path() matches on sample number, not filename pattern"]],
      widths=[1.15, 2.75, 2.5], size=8.5)

h2("10.3 Interface additions")
table(["Addition", "Detail"],
      [["extractor parameter", "process_pdf() selects the backend directly; previously reachable only by "
                               "constructing a Strategy by hand"],
       ["apply_rules parameter", "Path B conversion wired through process_pdf(), run_and_save() and the CLI "
                                 "as --apply-rules"],
       ["_write_outputs() helper", "Single code path writing flat and structured JSON, shared by process_pdf() "
                                   "and the batch runners"],
       ["WholeDocStrategy export", "Exposed from the package root for direct construction"],
       ["LightOnOCR bold signal", "Markdown heading markers in OCR output set is_bold and are stripped from "
                                  "the text, giving the classifier the emphasis cue the other extractors take "
                                  "from font data"],
       ["Four-stage output layout", "core/paths.py defines 1_extracted / 2_labeled / 3_structured / "
                                    "4_final. Stages 3 and 4 previously shared a directory with stage 2, "
                                    "distinguished only by a _structured filename suffix, and --apply-rules "
                                    "overwrote stage 3 in place so the converted and unconverted forms could "
                                    "not both exist. Both are now always written"],
       ["Extraction cache", "Stage 1 is keyed by extractor, so it is computed once and reused across models "
                            "rather than repeated per run"]],
      widths=[1.55, 4.85], size=8.5)

h2("10.4 Duplicated system prompt — resolved")
para("Until 5 August the prompt file contained the prompt twice, and the two copies contradicted each other "
     "on the boundary the models handled worst:")
table(["Label","First copy said","Second copy said"],
      [["section.title","Often starts with a letter prefix (A., B., C.)",
        "Does NOT include sub-items labelled A., B., C."],
       ["question.text","A specific question, instruction or prompt",
        "Often appears as a lettered sub-item (A., B., C.)"]],
      widths=[1.2, 2.6, 2.6], size=8.5)
para("Removing the duplicate cost roughly 16 points on pdfplumber, which was recorded at the time as "
     "unexplained. It is now explained: the deleted second copy was the one carrying the correct guidance. "
     "The surviving copy told the model that lettered sub-items are section titles, which is what section 4.2 "
     "corrects. Restoring that guidance recovered the loss and more.")
para("The duplication existed only in the working file and was never committed, so it does not appear in the "
     "repository history.", size=10, color=GREY)

# ── 11 ───────────────────────────────────────────────────────────────────
h1("11. Open Questions and Next Steps")
table(["Item","Description"],
      [["Decide the prompt  (first)",
        "The best llama3.1:8b wording is the worst gemma4:e4b wording — the same edit gained 0.6 points on one "
        "and cost 4.0 on the other (section 4.2). Either accept a per-model prompt, or optimise for gemma4:e4b "
        "and llama3.3:70b and drop the 8B. Every result in section 6 depends on which is chosen."],
       ["Re-run every model after any prompt change",
        "Prompt effects do not transfer between models, and the prompt is whitespace-significant — three blank "
        "lines moved cost 1.9 points of F1 with no wording change. A partial re-run produces figures that "
        "cannot be compared."],
       ["Investigate recurring data loss (new, 19 August)",
        "An entire model's pipeline output has vanished from disk without this session deleting it, four "
        "separate times across two days (13 and 19 August) — most recently gemma4:e4b's full output "
        "mid-session. Root cause not found; something outside this session appears to have write/delete "
        "access to the working tree. Needs a direct investigation, not another re-run to paper over it."],
       ["Consider dropping llama3.1:8b",
        "14 points behind gemma4:e4b at identical runtime, and unable to act on the prompt correction that "
        "both other models benefit from."],
       ["Expand the evaluation set",
        "Ten documents is few. question.text has 16-18 gold items, which is workable but thin, and one "
        "document can move a per-label score by tens of points."],
       ["Test a second labeling strategy",
        "Whole-doc has never been compared against anything. Page-by-page would be the simplest alternative "
        "and would also indicate behaviour on documents longer than those tested."],
       ["Fine-tuning (F01)",
        "LoRA fine-tune of a small model — deferred until the evaluation set can measure the result."]],
      widths=[1.85, 4.55])

out = Path("Report-doc/project_report.docx")
doc.save(out)
print(f"saved -> {out}  ({out.stat().st_size / 1024:.0f} KB)")
